import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

template = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\DDJJ\14 Declaración Jurada conocimiento Pliegos.docx'
output = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica\5 CV Hernán Hamra - Director de Proyecto.docx'

doc = Document(template)
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_para(p):
    elem = p._element
    pPr = elem.find('{%s}pPr' % ns)
    for child in list(elem):
        if child.tag != '{%s}pPr' % ns:
            elem.remove(child)
    if pPr is not None:
        numPr_elem = pPr.find('{%s}numPr' % ns)
        if numPr_elem is not None:
            pPr.remove(numPr_elem)

def set_text(p, text, bold=False, size=11):
    clear_para(p)
    elem = p._element
    new_run = etree.SubElement(elem, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def add_para(doc, after_elem, text, bold=False, size=11):
    new_p = etree.Element('{%s}p' % ns)
    after_elem.addnext(new_p)
    pPr = etree.SubElement(new_p, '{%s}pPr' % ns)
    spacing = etree.SubElement(pPr, '{%s}spacing' % ns)
    spacing.set('{%s}after' % ns, '40')
    spacing.set('{%s}line' % ns, '276')
    spacing.set('{%s}lineRule' % ns, 'auto')
    new_run = etree.SubElement(new_p, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return new_p

# --- Modify template ---

# P7: CV title
set_text(doc.paragraphs[7], 'CURRICULUM VITAE - Hernán Hamra', bold=True, size=14)

# Clear destinatario and saludo
set_text(doc.paragraphs[3], '')
set_text(doc.paragraphs[4], '')
set_text(doc.paragraphs[5], '')
set_text(doc.paragraphs[8], '')

# P9: Rol
set_text(doc.paragraphs[9], 'Rol en el Proyecto: Director de Proyecto', bold=True, size=11)

# P10: Clear
set_text(doc.paragraphs[10], '')
# P11: Clear
set_text(doc.paragraphs[11], '')

# Remove paragraphs 12 onwards
body = doc._element.body
paras = list(body.findall('{%s}p' % ns))
for i in range(12, len(paras)):
    body.remove(paras[i])

content = [
    ('', False, 11),
    ('DATOS PERSONALES', True, 12),
    ('Nombre completo: Hernán Hamra', False, 11),
    ('DNI: 23.505.172', False, 11),
    ('Cargo en la empresa: Apoderado - Software By Design S.A.', False, 11),
    ('Email: hernan@softwarebydesign.com.ar', False, 11),
    ('Teléfono: +54 11 3953-0850', False, 11),
    ('', False, 11),
    ('PERFIL PROFESIONAL', True, 12),
    ('Profesional con amplia experiencia en gestión de proyectos de tecnología e '
     'infraestructura IT, licitaciones públicas y desarrollo de soluciones tecnológicas. '
     'Responsable de la dirección comercial y operativa de Software By Design S.A., '
     'con trayectoria en la ejecución exitosa de proyectos para organismos del Estado '
     'Nacional, Provincial y Municipal, así como para el sector privado.', False, 10),
    ('', False, 11),
    ('EXPERIENCIA LABORAL', True, 12),
    ('', False, 11),
    ('Software By Design S.A. - Apoderado / Director de Proyectos', True, 11),
    ('2008 - Presente', False, 10),
    ('Dirección y coordinación integral de proyectos de infraestructura de redes, '
     'cableado estructurado, telefonía IP, seguridad electrónica (CCTV y control de acceso), '
     'UPS y sistemas de telecomunicaciones. Gestión de licitaciones públicas y privadas. '
     'Relación directa con comitentes y seguimiento de obra.', False, 10),
    ('', False, 11),
    ('Principales proyectos dirigidos:', False, 11),
    ('    - SBASE: UPS para trenes de subterráneos', False, 10),
    ('    - ADIF: Relevamiento de redes, cableado, datacenter, cámaras de seguridad', False, 10),
    ('    - Ministerio de Seguridad PBA: Redes WiFi (P12), ampliación WiFi, red LAN, UPS', False, 10),
    ('    - Presidencia de la Nación: Insumos de red', False, 10),
    ('    - Hospital de Malvinas Argentinas: Cableado estructurado', False, 10),
    ('    - Hospital Posadas: Cableado estructurado y networking', False, 10),
    ('    - Municipalidad de Bolívar: Seguridad urbana CCTV', False, 10),
    ('    - Hospital Garrahan: Storage Huawei, actualización Oracle BD', False, 10),
    ('', False, 11),
    ('COMPETENCIAS', True, 12),
    ('    - Dirección y gestión de proyectos de infraestructura IT', False, 11),
    ('    - Licitaciones públicas y privadas (preparación, presentación y ejecución)', False, 11),
    ('    - Coordinación de equipos técnicos multidisciplinarios', False, 11),
    ('    - Gestión contractual y administrativa de proyectos', False, 11),
    ('    - Cableado estructurado y redes LAN/WAN', False, 11),
    ('    - Telefonía IP y comunicaciones', False, 11),
    ('    - Seguridad electrónica (CCTV, control de acceso)', False, 11),
    ('    - Desarrollo de software y soluciones tecnológicas', False, 11),
    ('', False, 11),
]

# Fill P10, then add rest
set_text(doc.paragraphs[10], content[0][0], bold=content[0][1], size=content[0][2])

last_elem = doc.paragraphs[10]._element
for text, bold, size in content[1:]:
    last_elem = add_para(doc, last_elem, text, bold=bold, size=size)

# Remove old P11
old_p11 = body.findall('{%s}p' % ns)[11]
body.remove(old_p11)

doc.save(output)
print(f'Guardado: {output}')
print('OK')
