import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

template = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\DDJJ\14 Declaración Jurada conocimiento Pliegos.docx'
output = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica\4 CV Marcelo Hamra - Director Técnico.docx'

doc = Document(template)
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_para(p):
    elem = p._element
    pPr = elem.find('{%s}pPr' % ns)
    for child in list(elem):
        if child.tag != '{%s}pPr' % ns:
            elem.remove(child)
    if pPr is not None:
        numPr_elem = pPr.find('{%s}numPr' % ns)
        if numPr_elem is not None:
            pPr.remove(numPr_elem)

def set_text(p, text, bold=False, size=11):
    clear_para(p)
    elem = p._element
    new_run = etree.SubElement(elem, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def add_para(doc, after_elem, text, bold=False, size=11):
    new_p = etree.Element('{%s}p' % ns)
    after_elem.addnext(new_p)
    pPr = etree.SubElement(new_p, '{%s}pPr' % ns)
    spacing = etree.SubElement(pPr, '{%s}spacing' % ns)
    spacing.set('{%s}after' % ns, '40')
    spacing.set('{%s}line' % ns, '276')
    spacing.set('{%s}lineRule' % ns, 'auto')
    new_run = etree.SubElement(new_p, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return new_p

# --- Modify template ---

# Remove destinatario (P3-5) and saludo (P8) - not needed for CV
# P7: Change reference to CV title
set_text(doc.paragraphs[7], 'CURRICULUM VITAE - Marcelo Ariel Hamra', bold=True, size=14)

# P3-5: Clear destinatario
set_text(doc.paragraphs[3], '')
set_text(doc.paragraphs[4], '')
set_text(doc.paragraphs[5], '')

# P8: Clear saludo
set_text(doc.paragraphs[8], '')

# P9: Start CV content - Datos personales header
set_text(doc.paragraphs[9], 'Rol en el Proyecto: Director Técnico', bold=True, size=11)

# P10: Clear (hyperlink paragraph)
set_text(doc.paragraphs[10], '')

# P11: Clear cierre
set_text(doc.paragraphs[11], '')

# Remove paragraphs 12 onwards
body = doc._element.body
paras = list(body.findall('{%s}p' % ns))
for i in range(12, len(paras)):
    body.remove(paras[i])

# Build CV content after paragraph 9
content = [
    # After P9 (rol), P10 is empty, we'll use it
    ('', False, 11),
    ('DATOS PERSONALES', True, 12),
    ('Nombre completo: Marcelo Ariel Hamra', False, 11),
    ('DNI: 20.665.853', False, 11),
    ('CUIL: 20-20665853-4', False, 11),
    ('Cargo en la empresa: Presidente - Software By Design S.A.', False, 11),
    ('', False, 11),
    ('PERFIL PROFESIONAL', True, 12),
    ('Profesional con más de 15 años de experiencia en consultoría tecnológica, '
     'arquitectura de soluciones empresariales y gestión de proyectos de infraestructura IT. '
     'Amplia trayectoria en diseño e implementación de soluciones de redes, telecomunicaciones, '
     'seguridad electrónica y sistemas de información para organismos públicos y empresas privadas.', False, 10),
    ('', False, 11),
    ('EXPERIENCIA LABORAL', True, 12),
    ('', False, 11),
    ('Software By Design S.A. - Presidente y Director Técnico', True, 11),
    ('2008 - Presente', False, 10),
    ('Dirección técnica y comercial de la empresa. Diseño de soluciones de infraestructura '
     'de redes, cableado estructurado, telefonía IP, CCTV, control de acceso y UPS para '
     'organismos públicos (SBASE, ADIF, Ministerio de Seguridad PBA, Presidencia de la Nación, '
     'hospitales públicos) y sector privado. Gestión integral de proyectos desde la propuesta '
     'hasta la implementación y entrega.', False, 10),
    ('', False, 11),
    ('Hewlett Packard Enterprise - Consultor Senior', True, 11),
    ('12 años de trayectoria', False, 10),
    ('Consultoría en soluciones enterprise, arquitectura de sistemas, infraestructura IT '
     'y gestión de proyectos tecnológicos para grandes corporaciones y organismos gubernamentales.', False, 10),
    ('', False, 11),
    ('Oracle - Consultor', True, 11),
    ('Implementación y consultoría en bases de datos y soluciones empresariales Oracle.', False, 10),
    ('', False, 11),
    ('Senado de la Nación Argentina', True, 11),
    ('Gestión de infraestructura tecnológica y sistemas de información.', False, 10),
    ('', False, 11),
    ('Soluciones Y Proyectos Informáticos', True, 11),
    ('Consultoría y desarrollo de proyectos de tecnología.', False, 10),
    ('', False, 11),
    ('COMPETENCIAS TÉCNICAS', True, 12),
    ('    - Arquitectura de soluciones de infraestructura IT', False, 11),
    ('    - Diseño de redes LAN/WAN y cableado estructurado', False, 11),
    ('    - Telefonía IP y comunicaciones unificadas', False, 11),
    ('    - Sistemas de CCTV y seguridad electrónica', False, 11),
    ('    - Control de acceso', False, 11),
    ('    - Sistemas de energía ininterrumpida (UPS)', False, 11),
    ('    - Gestión de proyectos (Project Management)', False, 11),
    ('    - ITIL - Gestión de servicios IT', False, 11),
    ('    - Enterprise Software', False, 11),
    ('    - Sales Management', False, 11),
    ('', False, 11),
    ('PROYECTOS RELEVANTES EN SOFTWARE BY DESIGN S.A.', True, 12),
    ('    - SBASE: Suministro e instalación de equipamiento UPS para trenes', False, 10),
    ('    - ADIF: Relevamiento de redes, suministro de materiales de cableado, tablero datacenter, cámaras de seguridad', False, 10),
    ('    - Ministerio de Seguridad PBA: Redes WiFi, ampliación WiFi, red LAN, sistemas UPS', False, 10),
    ('    - Presidencia de la Nación: Insumos de red', False, 10),
    ('    - Hospital de Malvinas Argentinas: Cableado estructurado', False, 10),
    ('    - Municipalidad de Bolívar: CCTV seguridad urbana', False, 10),
    ('    - Hospital Posadas: Cableado estructurado y networking', False, 10),
    ('    - Hospital Garrahan: Storage Huawei, actualización Oracle BD', False, 10),
    ('', False, 11),
]

# Fill P10 with first content line, then add rest
set_text(doc.paragraphs[10], content[0][0], bold=content[0][1], size=content[0][2])

last_elem = doc.paragraphs[10]._element
for text, bold, size in content[1:]:
    last_elem = add_para(doc, last_elem, text, bold=bold, size=size)

# Remove old P11 (was cierre, now empty)
old_p11 = body.findall('{%s}p' % ns)[11]
body.remove(old_p11)

doc.save(output)
print(f'Guardado: {output}')
print('OK')
