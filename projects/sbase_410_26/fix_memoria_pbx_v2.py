import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

base = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica'
path = base + r'\varios\6 Memoria Técnica y Plan de Trabajo.docx'
doc = Document(path)

changes = 0

replacements = [
    # Fix "1,500 usuarios" variant (with comma)
    ('1,500 usuarios', '250 usuarios'),
    # Fix the duplicate model line
    ('Grandstream HT881/HT881', 'Grandstream HT881'),
    # Fix the incorrect spec line (HT881 has 8 FXO, not 4)
    ('HT881: 4 puertos FXO + 1 FXS / HT881: 8 puertos FXO +', 'HT881: 8 puertos FXO + 1 FXS, 2x GbE, T.38 fax,'),
    # Also catch simpler variant
    ('HT881: 4 puertos FXO', 'HT881: 8 puertos FXO'),
]

for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        orig = run.text
        for old, new in replacements:
            if old in run.text:
                run.text = run.text.replace(old, new)
        if run.text != orig:
            print(f'P[{i}]: {orig[:80]}')
            print(f'    -> {run.text[:80]}')
            changes += 1

for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    orig = run.text
                    for old, new in replacements:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                    if run.text != orig:
                        print(f'T[{t_idx}]R[{r_idx}]C[{c_idx}]: {orig[:80]}')
                        print(f'    -> {run.text[:80]}')
                        changes += 1

doc.save(path)
print(f'\nTotal cambios: {changes}')

if changes > 0:
    print('Guardado .docx OK - Regenerando PDF...')
    import subprocess, time
    subprocess.run(['taskkill', '/f', '/im', 'WINWORD.EXE'], capture_output=True)
    time.sleep(1)

    import win32com.client
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    try:
        wdoc = word.Documents.Open(path)
        pdf_path = base + r'\6 Memoria Técnica y Plan de Trabajo.pdf'
        wdoc.SaveAs(pdf_path, FileFormat=17)
        wdoc.Close()
        print(f'PDF regenerado OK')
    except Exception as e:
        print(f'Error PDF: {e}')
    finally:
        word.Quit()
else:
    print('No hay cambios adicionales.')
