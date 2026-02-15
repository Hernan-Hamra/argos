import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

template = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\DDJJ\14 Declaración Jurada conocimiento Pliegos.docx'
output = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica\3 Organigrama del Proyecto.docx'

doc = Document(template)
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_para(p):
    elem = p._element
    pPr = elem.find('{%s}pPr' % ns)
    for child in list(elem):
        if child.tag != '{%s}pPr' % ns:
            elem.remove(child)
    if pPr is not None:
        numPr_elem = pPr.find('{%s}numPr' % ns)
        if numPr_elem is not None:
            pPr.remove(numPr_elem)

def set_text(p, text, bold=False, size=11):
    clear_para(p)
    elem = p._element
    new_run = etree.SubElement(elem, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def add_para(doc, after_elem, text, bold=False, size=11):
    new_p = etree.Element('{%s}p' % ns)
    after_elem.addnext(new_p)
    pPr = etree.SubElement(new_p, '{%s}pPr' % ns)
    spacing = etree.SubElement(pPr, '{%s}spacing' % ns)
    spacing.set('{%s}after' % ns, '40')
    spacing.set('{%s}line' % ns, '276')
    spacing.set('{%s}lineRule' % ns, 'auto')
    new_run = etree.SubElement(new_p, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return new_p

# --- Modify template ---

# P7: Reference
set_text(doc.paragraphs[7], 'Ref: ORGANIGRAMA DEL PROYECTO - Licitación Privada N° 410/26', bold=True)

# P9: Intro
set_text(doc.paragraphs[9],
    'En cumplimiento con los requisitos establecidos en el Pliego de Bases y Condiciones y el '
    'Pliego de Especificaciones Técnicas de la Licitación Privada N° 410/26, se presenta a '
    'continuación el organigrama propuesto para la ejecución del proyecto "Implementación '
    'integral de infraestructura de redes y telecomunicaciones en edificio Balcarce N° 340".')

# P10: Clear (hyperlink)
set_text(doc.paragraphs[10], '')
# P11: Clear
set_text(doc.paragraphs[11], '')

# Remove paragraphs 12 onwards
body = doc._element.body
paras = list(body.findall('{%s}p' % ns))
for i in range(12, len(paras)):
    body.remove(paras[i])

# Build content
content = [
    ('', False, 11),
    ('ESTRUCTURA ORGANIZATIVA DEL PROYECTO', True, 13),
    ('', False, 11),
    ('1. Director de Proyecto', True, 11),
    ('Hernán Hamra - DNI 23.505.172', False, 11),
    ('Apoderado de Software By Design S.A.', False, 11),
    ('Responsable de la dirección general del proyecto, coordinación con SBASE, '
     'planificación, seguimiento de plazos, gestión contractual y administrativa. '
     'Punto de contacto principal con el comitente.', False, 10),
    ('', False, 11),
    ('2. Director Técnico', True, 11),
    ('Marcelo Ariel Hamra - DNI 20.665.853', False, 11),
    ('Presidente de Software By Design S.A.', False, 11),
    ('Responsable de la supervisión técnica integral del proyecto, definición de '
     'arquitectura de red, validación de soluciones tecnológicas, control de calidad '
     'y cumplimiento de especificaciones técnicas del pliego. Más de 15 años de '
     'experiencia en consultoría de infraestructura IT, arquitectura de soluciones '
     'y gestión de proyectos tecnológicos.', False, 10),
    ('', False, 11),
    ('3. Equipo Técnico de Instalación', True, 11),
    ('', False, 11),
    ('3.1. Esteban Agustín Marcuzzi - DNI 23.515.129', True, 11),
    ('Técnico Electrónico', False, 11),
    ('Técnico con amplia experiencia en instalaciones de redes de computación, '
     'mantenimiento industrial y sistemas electrónicos. Participación en proyectos '
     'de instalación de redes para ADIF y Comisarías del Ministerio de Seguridad PBA.', False, 10),
    ('', False, 11),
    ('3.2. José María Rodriguez - DNI 23.418.511', True, 11),
    ('Técnico Electromecánico / Especialista en Redes y Sistemas', False, 11),
    ('Consultor independiente en sistemas y redes de datos. Experiencia en cableado '
     'estructurado Cat5e/Cat6, fibra óptica, CCTV IP, y administración de redes '
     '(Cisco, Alcatel, Juniper, Fortinet). Ex Administrador de Redes del Ministerio '
     'de Seguridad PBA (sistema 911).', False, 10),
    ('', False, 11),
    ('3.3. Emanuel Andrés Delpino - DNI 41.432.997', True, 11),
    ('Técnico Instalador de Redes', False, 11),
    ('Técnico especializado en instalaciones de redes de computación. Participación '
     'en proyectos de instalación de infraestructura de red para Comisarías del '
     'Ministerio de Seguridad PBA.', False, 10),
    ('', False, 11),
    ('', False, 11),
    ('FUNCIONES DEL EQUIPO TÉCNICO', True, 13),
    ('', False, 11),
    ('El equipo técnico será responsable de:', False, 11),
    ('    - Instalación de cableado estructurado categoría 6A', False, 11),
    ('    - Tendido de fibra óptica', False, 11),
    ('    - Canalización y bandejas portacables', False, 11),
    ('    - Montaje de racks y patch panels', False, 11),
    ('    - Instalación de puntos de red y telefonía', False, 11),
    ('    - Configuración física de equipamiento activo (switches, APs, UPS)', False, 11),
    ('    - Certificación y etiquetado de puntos de red', False, 11),
    ('    - Pruebas de conectividad y puesta en marcha', False, 11),
    ('', False, 11),
    ('', False, 11),
    ('ESQUEMA JERÁRQUICO', True, 13),
    ('', False, 11),
]

# ASCII org chart - updated without Dnet
chart = [
    '              ┌───────────────────────────────────┐',
    '              │       Director de Proyecto         │',
    '              │          Hernán Hamra              │',
    '              └────────────────┬──────────────────┘',
    '                               │',
    '              ┌────────────────┴──────────────────┐',
    '              │        Director Técnico            │',
    '              │        Marcelo Hamra               │',
    '              └────────────────┬──────────────────┘',
    '                               │',
    '              ┌────────────────┴──────────────────┐',
    '              │     Equipo Técnico de Instalación  │',
    '              ├───────────────────────────────────┤',
    '              │  Esteban Marcuzzi - Téc. Electr.  │',
    '              │  José M. Rodriguez - Téc. Redes   │',
    '              │  Emanuel Delpino - Téc. Instalador │',
    '              └───────────────────────────────────┘',
]

for line in chart:
    content.append((line, False, 9))

# Closing
content.extend([
    ('', False, 11),
    ('', False, 11),
    ('Sin otro particular, saluda a ustedes atentamente,', False, 11),
    ('', False, 11),
    ('', False, 11),
    ('', False, 11),
    ('', False, 11),
    ('Hernán Hamra', False, 11),
    ('DNI:23.505.172 ', False, 11),
    ('Apoderado', False, 11),
    ('Software By Design S.A.', False, 11),
])

# Fill P10 with first content line, then add rest
set_text(doc.paragraphs[10], content[0][0], bold=content[0][1], size=content[0][2])

last_elem = doc.paragraphs[10]._element
for text, bold, size in content[1:]:
    last_elem = add_para(doc, last_elem, text, bold=bold, size=size)

# Remove old P11
old_p11 = body.findall('{%s}p' % ns)[11]
body.remove(old_p11)

doc.save(output)
print(f'Guardado: {output}')
print('OK')
