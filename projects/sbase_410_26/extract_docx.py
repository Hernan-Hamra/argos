import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

docx_path = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\VARIOS\PROPUESTA ALCANCE 3 FINAL\PROPUESTA COMERCIAL SBD- SBASE Conectividad y Telefonia Nro 70.712.docx'

doc = Document(docx_path)

print("="*80)
print("PARAGRAPHS")
print("="*80)
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[P{i:03d}] {text}")

print(f"\n{'='*80}")
print("TABLES")
print("='*80")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- TABLE {t_idx+1} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
        print(f"  Row {r_idx}: {' || '.join(cells)}")
