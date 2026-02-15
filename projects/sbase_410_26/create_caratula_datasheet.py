import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

base = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR'
template = base + r'\DDJJ\14 Declaración Jurada conocimiento Pliegos.docx'
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag != '{%s}sectPr' % ns:
            body.remove(child)

def add_para(doc, text, bold=False, size=11, align=None, space_after=6,
             space_before=2, font='Calibri', color=None):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = font
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        r_elem = run._element
        rPr = r_elem.find('{%s}rPr' % ns)
        if rPr is None:
            rPr = etree.SubElement(r_elem, '{%s}rPr' % ns)
        rFonts = rPr.find('{%s}rFonts' % ns)
        if rFonts is None:
            rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
        rFonts.set('{%s}ascii' % ns, font)
        rFonts.set('{%s}hAnsi' % ns, font)
        rFonts.set('{%s}cs' % ns, font)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run('_' * 60)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(27, 79, 114)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(12)

doc = Document(template)
clear_body(doc)

add_para(doc, '', space_after=30)
add_para(doc, '', space_after=30)

add_para(doc, 'LICITACIÓN PRIVADA N° 410/26', size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
         color=(27, 79, 114))

add_para(doc, 'SUBTERRÁNEOS DE BUENOS AIRES S.E.', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
         color=(27, 79, 114))

add_para(doc, 'Implementación integral de infraestructura de redes', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc, 'y telecomunicaciones en edificio Balcarce N° 340', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

add_separator(doc)

add_para(doc, 'DATASHEETS', size=28, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
         color=(27, 79, 114))

add_para(doc, 'Folletos y Especificaciones Técnicas', size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
         color=(93, 109, 126))

add_separator(doc)

add_para(doc, 'CONTENIDO:', size=11, bold=True, space_after=8,
         color=(27, 79, 114))

datasheets = [
    ('Furukawa', [
        'Cable GigaLan Cat6A U/UTP',
        'Conector Hembra Cat6A',
        'Patch Cord Cat6A',
        'Patch Panel Modular GigaLan Cat6',
        'Faceplate Modular',
        'Cable Fiber-LAN Indoor OM3',
        'DIO BW12 ODF',
    ]),
    ('Ubiquiti', [
        'UniFi Pro 48 PoE (USW-Pro-48-POE)',
        'UniFi Cloud Gateway Fiber (UCG-Fiber)',
        'UniFi U6 Long-Range (U6-LR)',
    ]),
    ('Grandstream', [
        'UCM6300A – Central IP 250 usuarios',
        'GRP2601P – Teléfono IP',
        'HT881 – Gateway FXO 8 puertos',
    ]),
    ('Panduit', [
        'Rack 42U 4 Postes',
        'Rack Net-Access',
    ]),
    ('Kaise', [
        'UPS 1-3 kVA Rack',
    ]),
]

for marca, items in datasheets:
    add_para(doc, marca, size=11, bold=True, space_after=3, space_before=6,
             color=(27, 79, 114))
    for item in items:
        p = add_para(doc, f'•  {item}', size=10, space_after=2, space_before=1)
        p.paragraph_format.left_indent = Cm(1.5)

add_para(doc, '', space_after=20)
add_separator(doc)
add_para(doc, 'SOFTWARE BY DESIGN S.A.', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
         color=(27, 79, 114))
add_para(doc, 'CUIT 30-70894532-0', size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
         color=(93, 109, 126))
add_para(doc, 'Febrero 2026', size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
         color=(93, 109, 126))

out_dir = base + r'\CARPETA B Documentación Técnica'
out_docx = out_dir + r'\varios\10 Carátula Datasheets.docx'
doc.save(out_docx)
print(f'Guardado: {out_docx}')

# Generar PDF
import subprocess, time
subprocess.run(['taskkill', '/f', '/im', 'WINWORD.EXE'], capture_output=True)
time.sleep(1)

import win32com.client
word = win32com.client.Dispatch('Word.Application')
word.Visible = False
try:
    wdoc = word.Documents.Open(out_docx)
    pdf_path = out_dir + r'\10 Carátula Datasheets.pdf'
    wdoc.SaveAs(pdf_path, FileFormat=17)
    wdoc.Close()
    print(f'PDF: {pdf_path}')
except Exception as e:
    print(f'Error PDF: {e}')
finally:
    word.Quit()
