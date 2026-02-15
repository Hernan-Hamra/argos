import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

base = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica'
path = base + r'\varios\6 Memoria Técnica y Plan de Trabajo.docx'
doc = Document(path)

changes = 0

for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        orig = run.text
        if '1500 usuarios' in run.text:
            run.text = run.text.replace('1500 usuarios', '250 usuarios')
        if '1.500 usuarios' in run.text:
            run.text = run.text.replace('1.500 usuarios', '250 usuarios')
        if '200 llamadas simultáneas' in run.text:
            run.text = run.text.replace('200 llamadas simultáneas', '50 llamadas simultáneas')
        if '200 llamadas simult.' in run.text:
            run.text = run.text.replace('200 llamadas simult.', '50 llamadas simult.')
        if 'HT841' in run.text:
            run.text = run.text.replace('HT841', 'HT881')
        if '4 FXO' in run.text and 'HT' not in run.text:
            run.text = run.text.replace('4 FXO', '8 FXO')
        if run.text != orig:
            print(f'P[{i}]: {orig[:60]} -> {run.text[:60]}')
            changes += 1

for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    orig = run.text
                    if '1500 usuarios' in run.text:
                        run.text = run.text.replace('1500 usuarios', '250 usuarios')
                    if '1.500 usuarios' in run.text:
                        run.text = run.text.replace('1.500 usuarios', '250 usuarios')
                    if '200 llamadas simultáneas' in run.text:
                        run.text = run.text.replace('200 llamadas simultáneas', '50 llamadas simultáneas')
                    if '200 llamadas simult.' in run.text:
                        run.text = run.text.replace('200 llamadas simult.', '50 llamadas simult.')
                    if 'Hasta 1500' in run.text:
                        run.text = run.text.replace('Hasta 1500', 'Hasta 250')
                    if 'HT841' in run.text:
                        run.text = run.text.replace('HT841', 'HT881')
                    if '4 FXO' in run.text:
                        run.text = run.text.replace('4 FXO', '8 FXO')
                    if run.text != orig:
                        print(f'T[{t_idx}]R[{r_idx}]C[{c_idx}]: {orig[:60]} -> {run.text[:60]}')
                        changes += 1

doc.save(path)
print(f'\nTotal cambios: {changes}')
print('Guardado .docx OK')

# Regenerar PDF
import subprocess, time
subprocess.run(['taskkill', '/f', '/im', 'WINWORD.EXE'], capture_output=True)
time.sleep(1)

import win32com.client
word = win32com.client.Dispatch('Word.Application')
word.Visible = False
try:
    wdoc = word.Documents.Open(path)
    pdf_path = base + r'\6 Memoria Técnica y Plan de Trabajo.pdf'
    wdoc.SaveAs(pdf_path, FileFormat=17)
    wdoc.Close()
    print(f'PDF regenerado: {pdf_path}')
except Exception as e:
    print(f'Error PDF: {e}')
finally:
    word.Quit()
