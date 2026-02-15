import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

template = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\DDJJ\14 Declaración Jurada conocimiento Pliegos.docx'
output = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica\3 Organigrama del Proyecto.docx'

doc = Document(template)
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_and_set(p, text, bold=False, size=11, alignment=None):
    elem = p._element
    pPr = elem.find('{%s}pPr' % ns)
    for child in list(elem):
        if child.tag != '{%s}pPr' % ns:
            elem.remove(child)
    if pPr is not None:
        numPr_elem = pPr.find('{%s}numPr' % ns)
        if numPr_elem is not None:
            pPr.remove(numPr_elem)
        # Handle hyperlinks at pPr level
    # Also remove hyperlinks at paragraph level
    for hyp in list(elem.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')):
        elem.remove(hyp)

    new_run = etree.SubElement(elem, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    if alignment is not None:
        if pPr is None:
            pPr = etree.SubElement(elem, '{%s}pPr' % ns)
            elem.insert(0, pPr)
        jc = pPr.find('{%s}jc' % ns)
        if jc is None:
            jc = etree.SubElement(pPr, '{%s}jc' % ns)
        jc.set('{%s}val' % ns, alignment)

def add_para(doc, after_elem, text, bold=False, size=11):
    new_p = etree.Element('{%s}p' % ns)
    after_elem.addnext(new_p)
    pPr = etree.SubElement(new_p, '{%s}pPr' % ns)
    spacing = etree.SubElement(pPr, '{%s}spacing' % ns)
    spacing.set('{%s}after' % ns, '40')
    spacing.set('{%s}line' % ns, '276')
    spacing.set('{%s}lineRule' % ns, 'auto')
    new_run = etree.SubElement(new_p, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return new_p

# --- Modify template ---

# P7: Change reference
clear_and_set(doc.paragraphs[7], 'Ref: ORGANIGRAMA DEL PROYECTO - Licitación Privada N° 410/26', bold=True)

# P9: Replace intro
clear_and_set(doc.paragraphs[9],
    'En cumplimiento con los requisitos establecidos en el Pliego de Bases y Condiciones y el '
    'Pliego de Especificaciones Técnicas de la Licitación Privada N° 410/26, se presenta a '
    'continuación el organigrama propuesto para la ejecución del proyecto "Implementación '
    'integral de infraestructura de redes y telecomunicaciones en edificio Balcarce N° 340".')

# P10: Clear (had hyperlink)
clear_and_set(doc.paragraphs[10], '')

# P11: Clear cierre (will re-add later)
clear_and_set(doc.paragraphs[11], '')

# Remove paragraphs 12 onwards (blanks + signature) - we'll rebuild
body = doc._element.body
paras = list(body.findall('{%s}p' % ns))
for i in range(12, len(paras)):
    body.remove(paras[i])

# Now add content after paragraph 10 (which is now empty)
# Build content lines: (text, bold, size)
content = [
    ('', False, 11),
    ('ESTRUCTURA ORGANIZATIVA DEL PROYECTO', True, 13),
    ('', False, 11),
    ('1. Director de Proyecto', True, 11),
    ('Hernán Hamra - DNI 23.505.172', False, 11),
    ('Apoderado de Software By Design S.A.', False, 11),
    ('Responsable de la dirección general del proyecto, coordinación con SBASE, '
     'planificación, seguimiento de plazos, gestión contractual y administrativa. '
     'Punto de contacto principal con el comitente.', False, 10),
    ('', False, 11),
    ('2. Director Técnico', True, 11),
    ('Marcelo Ariel Hamra - DNI 20.665.853', False, 11),
    ('Presidente de Software By Design S.A.', False, 11),
    ('Responsable de la supervisión técnica integral del proyecto, definición de '
     'arquitectura de red, validación de soluciones tecnológicas, control de calidad '
     'y cumplimiento de especificaciones técnicas del pliego. Más de 15 años de '
     'experiencia en consultoría de infraestructura IT, arquitectura de soluciones '
     'y gestión de proyectos tecnológicos.', False, 10),
    ('', False, 11),
    ('3. Técnico de Obra / Jefe de Instalación', True, 11),
    ('Leonardo Rodriguez - Dnet', False, 11),
    ('Responsable de la ejecución en campo: instalación de cableado estructurado, '
     'tendido de fibra óptica, montaje de racks y canalizaciones, instalación de '
     'puntos de red y telefonía, configuración física de equipamiento activo '
     '(switches, APs, UPS). Coordinación directa del equipo de instaladores.', False, 10),
    ('', False, 11),
    ('4. Equipo Técnico de Instalación', True, 11),
    ('Personal técnico de Dnet especializado en:', False, 11),
    ('    - Cableado estructurado categoría 6A', False, 11),
    ('    - Fibra óptica', False, 11),
    ('    - Canalización y bandejas portacables', False, 11),
    ('    - Montaje de racks y patch panels', False, 11),
    ('    - Certificación y etiquetado de puntos de red', False, 11),
    ('', False, 11),
    ('', False, 11),
    ('ESQUEMA JERÁRQUICO', True, 13),
    ('', False, 11),
]

# ASCII org chart lines
chart = [
    '                    ┌───────────────────────────────┐',
    '                    │     Director de Proyecto       │',
    '                    │        Hernán Hamra            │',
    '                    └──────────────┬────────────────┘',
    '                                   │',
    '                    ┌──────────────┴────────────────┐',
    '                    │      Director Técnico          │',
    '                    │      Marcelo Hamra             │',
    '                    └──────────────┬────────────────┘',
    '                                   │',
    '                    ┌──────────────┴────────────────┐',
    '                    │     Técnico de Obra            │',
    '                    │   Leonardo Rodriguez (Dnet)    │',
    '                    └──────────────┬────────────────┘',
    '                                   │',
    '                    ┌──────────────┴────────────────┐',
    '                    │   Equipo de Instalación        │',
    '                    │    Personal Técnico Dnet       │',
    '                    └───────────────────────────────┘',
]

for line in chart:
    content.append((line, False, 9))

# Closing
content.extend([
    ('', False, 11),
    ('', False, 11),
    ('Sin otro particular, saluda a ustedes atentamente,', False, 11),
    ('', False, 11),
    ('', False, 11),
    ('', False, 11),
    ('', False, 11),
    ('Hernán Hamra', False, 11),
    ('DNI:23.505.172 ', False, 11),
    ('Apoderado', False, 11),
    ('Software By Design S.A.', False, 11),
])

# Insert content starting after paragraph 10
last_elem = doc.paragraphs[10]._element
for text, bold, size in content:
    last_elem = add_para(doc, last_elem, text, bold=bold, size=size)

# Remove the now-empty paragraph 11 (old cierre)
old_p11 = body.findall('{%s}p' % ns)[11]
body.remove(old_p11)

doc.save(output)
print(f'Guardado: {output}')
print('OK')
