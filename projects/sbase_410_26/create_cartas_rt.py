import sys
sys.stdout.reconfigure(encoding='utf-8')
import shutil
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

base = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR'
carpeta_b = base + r'\CARPETA B Documentación Técnica'
template = base + r'\DDJJ\14 Declaración Jurada conocimiento Pliegos.docx'

ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag != '{%s}sectPr' % ns:
            body.remove(child)

def add_para(doc, text, bold=False, size=11, align=None, space_after=6, space_before=2, font='Calibri'):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = font
        run.bold = bold
        r_elem = run._element
        rPr = r_elem.find('{%s}rPr' % ns)
        if rPr is None:
            rPr = etree.SubElement(r_elem, '{%s}rPr' % ns)
        rFonts = rPr.find('{%s}rFonts' % ns)
        if rFonts is None:
            rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
        rFonts.set('{%s}ascii' % ns, font)
        rFonts.set('{%s}hAnsi' % ns, font)
        rFonts.set('{%s}cs' % ns, font)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

# ============================================================
# 1. CARTA DE DESIGNACIÓN
# ============================================================
print("=== Creando Carta de Designación ===")
doc1 = Document(template)
clear_body(doc1)

add_para(doc1, '', space_after=2)
add_para(doc1, 'Buenos Aires, 18 de febrero de 2026', size=11,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

add_para(doc1, 'Señores', size=11, bold=True, space_after=2)
add_para(doc1, 'SUBTERRÁNEOS DE BUENOS AIRES S.E.', size=11, bold=True, space_after=2)
add_para(doc1, 'Presente', size=11, space_after=12)

add_para(doc1, 'Ref.: Licitación Privada N° 410/26 – Implementación integral de infraestructura '
         'de redes y telecomunicaciones en edificio Balcarce N° 340', size=11, bold=True, space_after=2)
add_para(doc1, 'Designación de Representante Técnico', size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(doc1, 'De nuestra consideración:', size=11, space_after=8)

add_para(doc1, 'Por medio de la presente, en mi carácter de representante legal de '
         'SOFTWARE BY DESIGN S.A. (CUIT 30-70894532-0), comunico a ustedes que la empresa '
         'designa al Ing. Marcelo Ariel Hamra, DNI 20.665.853, Ingeniero en Sistemas egresado '
         'de la Universidad Tecnológica Nacional (UTN), como Representante Técnico para la '
         'ejecución de los trabajos correspondientes a la Licitación Privada N° 410/26.',
         size=11, space_after=8)

add_para(doc1, 'El profesional designado se desempeña como Presidente y Director Técnico de '
         'Software By Design S.A., contando con más de 15 años de experiencia en consultoría '
         'tecnológica, arquitectura de soluciones de infraestructura IT, diseño de redes LAN/WAN, '
         'cableado estructurado, telefonía IP y gestión de proyectos tecnológicos.',
         size=11, space_after=8)

add_para(doc1, 'El Ing. Hamra ha dirigido técnicamente proyectos de similar envergadura y '
         'complejidad para organismos como SBASE, ADIF, Ministerio de Seguridad de la Provincia '
         'de Buenos Aires, Hospital Posadas, Hospital Garrahan, Presidencia de la Nación, entre otros, '
         'acreditando amplia experiencia en la posición conforme lo requerido en el artículo 5.12 '
         'del Pliego Único de Bases y Condiciones.',
         size=11, space_after=8)

add_para(doc1, 'Sin otro particular, saludo a ustedes muy atentamente.',
         size=11, space_after=24)

add_para(doc1, '', space_after=24)
add_para(doc1, '_________________________________', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc1, 'Hernán Hamra', size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc1, 'Representante Legal', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc1, 'Software By Design S.A.', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

out1 = carpeta_b + r'\4a Carta de Designación Representante Técnico.docx'
doc1.save(out1)
print(f'Guardado: {out1}')

# ============================================================
# 2. CARTA DE ACEPTACIÓN
# ============================================================
print("\n=== Creando Carta de Aceptación ===")
doc2 = Document(template)
clear_body(doc2)

add_para(doc2, '', space_after=2)
add_para(doc2, 'Buenos Aires, 18 de febrero de 2026', size=11,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

add_para(doc2, 'Señores', size=11, bold=True, space_after=2)
add_para(doc2, 'SUBTERRÁNEOS DE BUENOS AIRES S.E.', size=11, bold=True, space_after=2)
add_para(doc2, 'Presente', size=11, space_after=12)

add_para(doc2, 'Ref.: Licitación Privada N° 410/26 – Implementación integral de infraestructura '
         'de redes y telecomunicaciones en edificio Balcarce N° 340', size=11, bold=True, space_after=2)
add_para(doc2, 'Aceptación de Designación como Representante Técnico', size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(doc2, 'De mi consideración:', size=11, space_after=8)

add_para(doc2, 'Por medio de la presente, yo, Marcelo Ariel Hamra, DNI 20.665.853, '
         'Ingeniero en Sistemas egresado de la Universidad Tecnológica Nacional (UTN), '
         'manifiesto que acepto la designación como Representante Técnico de la empresa '
         'SOFTWARE BY DESIGN S.A. (CUIT 30-70894532-0) para la ejecución de los trabajos '
         'correspondientes a la Licitación Privada N° 410/26 – "Implementación integral de '
         'infraestructura de redes y telecomunicaciones en edificio Balcarce N° 340".',
         size=11, space_after=8)

add_para(doc2, 'Declaro conocer la totalidad de la documentación licitatoria, incluyendo el '
         'Pliego Único de Bases y Condiciones, el Pliego de Especificaciones Técnicas, sus '
         'Anexos y Circulares, y me comprometo a cumplir con las obligaciones inherentes al '
         'cargo de Representante Técnico conforme lo establecido en el artículo 5.12 y 8.3 '
         'del Pliego Único de Bases y Condiciones.',
         size=11, space_after=8)

add_para(doc2, 'Asimismo, declaro contar con la experiencia y competencias técnicas necesarias '
         'para desempeñar dicha función, habiendo participado en la dirección técnica de '
         'múltiples proyectos de infraestructura de redes y telecomunicaciones de similar '
         'envergadura y complejidad.',
         size=11, space_after=8)

add_para(doc2, 'Sin otro particular, saludo a ustedes muy atentamente.',
         size=11, space_after=24)

add_para(doc2, '', space_after=24)
add_para(doc2, '_________________________________', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc2, 'Ing. Marcelo Ariel Hamra', size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc2, 'DNI 20.665.853', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc2, 'Ingeniero en Sistemas - UTN', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

out2 = carpeta_b + r'\4b Carta de Aceptación Representante Técnico.docx'
doc2.save(out2)
print(f'Guardado: {out2}')

# ============================================================
# 3. ACTUALIZAR CV MARCELO - agregar Ingeniero en Sistemas UTN
# ============================================================
print("\n=== Actualizando CV Marcelo Hamra ===")
cv_path = carpeta_b + r'\4 CV Marcelo Hamra - Director Técnico.docx'
doc3 = Document(cv_path)

for i, p in enumerate(doc3.paragraphs):
    text = p.text.strip()
    # Update title line
    if text == 'CURRICULUM VITAE - Marcelo Ariel Hamra':
        for run in p.runs:
            if 'Marcelo' in run.text:
                run.text = 'CURRICULUM VITAE - Ing. Marcelo Ariel Hamra'
                break
        print(f'  [{i}] Updated title')

    # Add Título after DNI/CUIL lines - find the Cargo line and add before PERFIL
    if text == 'Cargo en la empresa: Presidente - Software By Design S.A.':
        # Add título universitario after this paragraph
        # We need to insert a new paragraph after this one
        elem = p._element
        new_p = etree.Element('{%s}p' % ns)
        elem.addnext(new_p)
        new_run = etree.SubElement(new_p, '{%s}r' % ns)
        rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
        rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
        rFonts.set('{%s}ascii' % ns, 'Calibri')
        rFonts.set('{%s}hAnsi' % ns, 'Calibri')
        sz = etree.SubElement(rPr, '{%s}sz' % ns)
        sz.set('{%s}val' % ns, '22')
        szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
        szCs.set('{%s}val' % ns, '22')
        t = etree.SubElement(new_run, '{%s}t' % ns)
        t.text = 'Título: Ingeniero en Sistemas - Universidad Tecnológica Nacional (UTN)'
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        # Add spacing
        pPr = etree.SubElement(new_p, '{%s}pPr' % ns)
        spacing = etree.SubElement(pPr, '{%s}spacing' % ns)
        spacing.set('{%s}after' % ns, '80')
        print(f'  [{i}] Added Título after Cargo')

    # Update PERFIL text to mention Ingeniero
    if 'Profesional con más de 15 años' in text:
        for run in p.runs:
            if 'Profesional con' in run.text:
                run.text = run.text.replace('Profesional con', 'Ingeniero en Sistemas (UTN) con', 1)
                print(f'  [{i}] Updated perfil text')
                break

    # Update Rol line
    if text == 'Rol en el Proyecto: Director Técnico':
        for run in p.runs:
            if 'Director Técnico' in run.text:
                run.text = 'Rol en el Proyecto: Director Técnico / Representante Técnico'
                print(f'  [{i}] Updated rol')
                break

doc3.save(cv_path)
print(f'Guardado: {cv_path}')

# ============================================================
# 4. ACTUALIZAR ORGANIGRAMA - Marcelo con título
# ============================================================
print("\n=== Actualizando Organigrama ===")
org_path = carpeta_b + r'\3 Organigrama del Proyecto.docx'
doc4 = Document(org_path)

for i, p in enumerate(doc4.paragraphs):
    text = p.text.strip()
    if 'Marcelo Ariel Hamra' in text and 'DNI' in text:
        for run in p.runs:
            if 'Marcelo' in run.text:
                run.text = run.text.replace('Marcelo Ariel Hamra', 'Ing. Marcelo Ariel Hamra')
                print(f'  [{i}] Updated: {run.text[:80]}')
                break
    # Also check the description paragraph after "2. Director Técnico"
    if 'Marcelo Ariel Hamra' in text and 'DNI' not in text and '│' not in text:
        for run in p.runs:
            if 'Marcelo' in run.text and 'Ing.' not in run.text:
                run.text = run.text.replace('Marcelo Ariel Hamra', 'Ing. Marcelo Ariel Hamra')
                print(f'  [{i}] Updated: {run.text[:80]}')
                break

doc4.save(org_path)
print(f'Guardado: {org_path}')

# ============================================================
# 5. ACTUALIZAR MEMORIA TÉCNICA - Marcelo con título
# ============================================================
print("\n=== Actualizando Memoria Técnica ===")
mem_path = carpeta_b + r'\6 Memoria Técnica y Plan de Trabajo.docx'
doc5 = Document(mem_path)

for i, p in enumerate(doc5.paragraphs):
    text = p.text.strip()
    if 'Marcelo Ariel Hamra' in text and 'Ing.' not in text:
        for run in p.runs:
            if 'Marcelo' in run.text and 'Ing.' not in run.text:
                run.text = run.text.replace('Marcelo Ariel Hamra', 'Ing. Marcelo Ariel Hamra')
                print(f'  Paragraph [{i}] Updated: {run.text[:80]}')
                break

# Update tables
for t_idx, table in enumerate(doc5.tables):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if 'Marcelo Ariel Hamra' in cell.text and 'Ing.' not in cell.text:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if 'Marcelo' in run.text and 'Ing.' not in run.text:
                            run.text = run.text.replace('Marcelo Ariel Hamra', 'Ing. Marcelo Ariel Hamra')
                            print(f'  Table[{t_idx}] R[{r_idx}] C[{c_idx}] Updated')

doc5.save(mem_path)
print(f'Guardado: {mem_path}')

print("\n=== TODO LISTO ===")
