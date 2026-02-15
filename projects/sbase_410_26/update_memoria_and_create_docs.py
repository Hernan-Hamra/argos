import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

base = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR'
carpeta_b = base + r'\CARPETA B Documentación Técnica'
template = base + r'\DDJJ\14 Declaración Jurada conocimiento Pliegos.docx'
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag != '{%s}sectPr' % ns:
            body.remove(child)

def add_para(doc, text, bold=False, size=11, align=None, space_after=6, space_before=2, font='Calibri'):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = font
        run.bold = bold
        r_elem = run._element
        rPr = r_elem.find('{%s}rPr' % ns)
        if rPr is None:
            rPr = etree.SubElement(r_elem, '{%s}rPr' % ns)
        rFonts = rPr.find('{%s}rFonts' % ns)
        if rFonts is None:
            rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
        rFonts.set('{%s}ascii' % ns, font)
        rFonts.set('{%s}hAnsi' % ns, font)
        rFonts.set('{%s}cs' % ns, font)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def set_cell_text(cell, text, bold=False, size=9):
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            p.runs[0].bold = bold
            p.runs[0].font.size = Pt(size)
        else:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = 'Calibri'

# ============================================================
# 1. ACTUALIZAR MEMORIA TÉCNICA
# ============================================================
print("=== 1. Actualizando Memoria Técnica ===")
mem_path = carpeta_b + r'\6 Memoria Técnica y Plan de Trabajo.docx'
doc = Document(mem_path)

# Update equipment table (Table 0)
table = doc.tables[0]

# R[8]: Rack -> Panduit
set_cell_text(table.rows[8].cells[1], 'Panduit Net-Access\nRack de piso 42U')
set_cell_text(table.rows[8].cells[2], '42U, 19", puertas con cerradura, bandejas, ventilación, gestión de cables')
print('  T[0] R[8] Rack -> Panduit')

# R[9]: Patch Panel -> Furukawa
set_cell_text(table.rows[9].cells[1], 'Furukawa GigaLan\nPatch Panel Modular 24p')
set_cell_text(table.rows[9].cells[2], 'Cat6A, 24 puertos, modular, montaje rack 19", T568A/B')
print('  T[0] R[9] Patch Panel -> Furukawa')

# R[10]: Cable UTP Cat6A -> Furukawa
set_cell_text(table.rows[10].cells[1], 'Furukawa GigaLan Cat.6A\nU/UTP 23AWGx4P LSZH')
set_cell_text(table.rows[10].cells[2], 'UTP, 500 MHz, 10GBASE-T, LSZH, cobre 23AWG, Bobina 305m')
print('  T[0] R[10] Cable UTP -> Furukawa')

# R[11]: Fibra óptica -> Furukawa
set_cell_text(table.rows[11].cells[1], 'Furukawa Fiber-LAN Indoor\nOM3 Multimodo')
set_cell_text(table.rows[11].cells[2], '6 hilos 50/125, 10GBASE-SR, LSZH, interior')
print('  T[0] R[11] Fibra OM3 -> Furukawa')

# R[12]: Patch cords -> Furukawa
set_cell_text(table.rows[12].cells[1], 'Furukawa GigaLan Cat.6A\nPatch Cord U/UTP')
set_cell_text(table.rows[12].cells[2], 'Cat6A certificados, LSZH, distintas longitudes')
print('  T[0] R[12] Patch cords -> Furukawa')

# Update paragraphs with brand references
updates_done = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text

    # Paragraph 59: cable description
    if 'tendido de cableado estructurado categoría 6A (UTP)' in text:
        for run in p.runs:
            if 'tendido de cableado' in run.text:
                run.text = run.text.replace(
                    'tendido de cableado estructurado categoría 6A (UTP)',
                    'tendido de cableado estructurado Furukawa GigaLan categoría 6A (U/UTP)'
                )
                updates_done += 1
                break

    # Paragraph about racks
    if 'racks de piso de 42 unidades de rack (42U)' in text:
        for run in p.runs:
            if 'racks de piso' in run.text:
                run.text = run.text.replace(
                    'racks de piso de 42 unidades de rack (42U)',
                    'racks de piso Panduit Net-Access de 42 unidades de rack (42U)'
                )
                updates_done += 1
                break

    # Patch panels reference
    if 'Patch panels Cat6A de 24 puertos' in text:
        for run in p.runs:
            if 'Patch panels Cat6A' in run.text:
                run.text = run.text.replace(
                    'Patch panels Cat6A de 24 puertos',
                    'Patch panels Furukawa GigaLan Cat6A de 24 puertos'
                )
                updates_done += 1
                break

    # Fibra óptica reference
    if 'cables de fibra óptica multimodo OM3' in text:
        for run in p.runs:
            if 'cables de fibra óptica multimodo OM3' in run.text:
                run.text = run.text.replace(
                    'cables de fibra óptica multimodo OM3',
                    'cables de fibra óptica Furukawa Fiber-LAN multimodo OM3'
                )
                updates_done += 1
                break

    # Cable Cat6A in phases
    if 'Tendido de cables UTP Cat6A' in text:
        for run in p.runs:
            if 'Tendido de cables UTP Cat6A' in run.text:
                run.text = run.text.replace(
                    'Tendido de cables UTP Cat6A',
                    'Tendido de cables Furukawa GigaLan Cat6A U/UTP'
                )
                updates_done += 1
                break

    # Jacks reference
    if 'jacks RJ45 Cat6A en face plates' in text:
        for run in p.runs:
            if 'jacks RJ45 Cat6A' in run.text:
                run.text = run.text.replace(
                    'jacks RJ45 Cat6A en face plates',
                    'conectores hembra Furukawa Cat6A en face plates'
                )
                updates_done += 1
                break

    # Patch panels in phases
    if 'patch panels Cat6A de 24 puertos en racks' in text:
        for run in p.runs:
            if 'patch panels Cat6A' in run.text:
                run.text = run.text.replace(
                    'patch panels Cat6A de 24 puertos en racks',
                    'patch panels Furukawa GigaLan Cat6A de 24 puertos en racks'
                )
                updates_done += 1
                break

    # Montaje de racks in phases
    if 'Montaje de racks de 42U' in text:
        for run in p.runs:
            if 'Montaje de racks de 42U' in run.text:
                run.text = run.text.replace(
                    'Montaje de racks de 42U',
                    'Montaje de racks Panduit de 42U'
                )
                updates_done += 1
                break

    # Patch cords reference
    if 'patch cords en patch panels' in text:
        for run in p.runs:
            if 'patch cords en patch panels' in run.text:
                run.text = run.text.replace(
                    'patch cords en patch panels',
                    'patch cords Furukawa Cat6A en patch panels'
                )
                updates_done += 1
                break

print(f'  Paragraphs updated: {updates_done}')

doc.save(mem_path)
print(f'  Guardado: {mem_path}')

# ============================================================
# 2. CREAR LISTA DE PROVEEDORES (Item 6 Carpeta B)
# ============================================================
print("\n=== 2. Creando Lista de Proveedores ===")
doc2 = Document(template)
clear_body(doc2)

add_para(doc2, '', space_after=2)
add_para(doc2, 'LISTA DE PROVEEDORES PROPUESTOS', size=14, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc2, 'Licitación Privada N° 410/26 - SBASE', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc2, 'Implementación integral de infraestructura de redes y telecomunicaciones', size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(doc2, 'A continuación se detalla la lista de proveedores propuestos para los materiales '
         'y equipamiento del proyecto, con indicación de marca, origen y representación local. '
         'Se adjuntan los folletos técnicos (datasheets) correspondientes en la carpeta DATASHEET.',
         size=10, space_after=12)

# Create table
table2 = doc2.add_table(rows=1, cols=5)
# Header
headers = ['Rubro', 'Marca / Fabricante', 'Origen', 'Representante Local', 'Productos']
for j, h in enumerate(headers):
    cell = table2.rows[0].cells[j]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'

# Data rows
proveedores = [
    ('Cableado estructurado\nCat6A (cobre)', 'Furukawa Electric\n(Línea GigaLan)', 'Brasil / Japón',
     'Furukawa Electric LatAm\nwww.furukawalatam.com',
     'Cable UTP Cat6A, Patch Panels, Conectores hembra, Patch Cords, Face Plates'),
    ('Fibra óptica\nmultimodo OM3', 'Furukawa Electric\n(Línea Fiber-LAN)', 'Brasil / Japón',
     'Furukawa Electric LatAm\nwww.furukawalatam.com',
     'Cable FO OM3 6 hilos, Distribuidor óptico (ODF), Patch Cords ópticos'),
    ('Switching y\nAccess Points', 'Ubiquiti Inc.\n(Línea UniFi)', 'Estados Unidos',
     'Ubiquiti Inc.\nui.com',
     'Switches UniFi Pro 48 PoE, Access Points U6-LR, Gateway UCG-Fiber'),
    ('Telefonía IP', 'Grandstream\nNetworks', 'Estados Unidos',
     'Grandstream Networks\nwww.grandstream.com',
     'Central PBX UCM6300A, Teléfonos GRP2601P, Gateway FXO HT841'),
    ('Energía\nininterrumpida', 'Kaise', 'China',
     'Kaise Argentina\nwww.kaiseups.com',
     'UPS Online Monofásico 1-3 kVA Rack'),
    ('Racks y gabinetes', 'Panduit', 'Estados Unidos',
     'Panduit Argentina\nwww.panduit.com/latam',
     'Racks de piso 42U Net-Access, accesorios de gestión de cables'),
]

for prov in proveedores:
    row = table2.add_row()
    for j, val in enumerate(prov):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

# Set column widths
from docx.shared import Inches
widths = [Cm(2.5), Cm(3.0), Cm(2.0), Cm(3.5), Cm(5.5)]
for row in table2.rows:
    for j, w in enumerate(widths):
        row.cells[j].width = w

# Format table borders
tbl = table2._tbl
tblPr = tbl.find('{%s}tblPr' % ns)
if tblPr is None:
    tblPr = etree.SubElement(tbl, '{%s}tblPr' % ns)
tblBorders = etree.SubElement(tblPr, '{%s}tblBorders' % ns)
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border = etree.SubElement(tblBorders, '{%s}%s' % (ns, border_name))
    border.set('{%s}val' % ns, 'single')
    border.set('{%s}sz' % ns, '4')
    border.set('{%s}space' % ns, '0')
    border.set('{%s}color' % ns, '000000')

# Header row shading
for cell in table2.rows[0].cells:
    tc = cell._element
    tcPr = tc.find('{%s}tcPr' % ns)
    if tcPr is None:
        tcPr = etree.SubElement(tc, '{%s}tcPr' % ns)
    shd = etree.SubElement(tcPr, '{%s}shd' % ns)
    shd.set('{%s}val' % ns, 'clear')
    shd.set('{%s}color' % ns, 'auto')
    shd.set('{%s}fill' % ns, '1B4F72')
    # White text
    for p in cell.paragraphs:
        for run in p.runs:
            r_elem = run._element
            rPr = r_elem.find('{%s}rPr' % ns)
            if rPr is None:
                rPr = etree.SubElement(r_elem, '{%s}rPr' % ns)
            color = etree.SubElement(rPr, '{%s}color' % ns)
            color.set('{%s}val' % ns, 'FFFFFF')

add_para(doc2, '', space_after=12)
add_para(doc2, 'Nota: Se adjuntan folletos explicativos y fichas técnicas (datasheets) de todos los '
         'productos en la carpeta DATASHEET que acompaña la presente documentación.',
         size=10, space_after=8)

add_para(doc2, 'Todos los proveedores propuestos cuentan con representación y soporte técnico '
         'en la República Argentina, garantizando disponibilidad de stock, asistencia técnica '
         'y respaldo de garantía para los productos ofrecidos.',
         size=10, space_after=12)

add_para(doc2, '', space_after=24)
add_para(doc2, '_________________________________', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc2, 'Hernán Hamra', size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc2, 'Representante Legal', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc2, 'Software By Design S.A.', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

out2 = carpeta_b + r'\7 Lista de Proveedores.docx'
doc2.save(out2)
print(f'  Guardado: {out2}')

# ============================================================
# 3. CREAR DDJJ PLAZO DE GARANTÍA (Item 9 Carpeta B)
# ============================================================
print("\n=== 3. Creando DDJJ Plazo de Garantía ===")
doc3 = Document(template)
clear_body(doc3)

add_para(doc3, '', space_after=2)
add_para(doc3, 'Buenos Aires, 18 de febrero de 2026', size=11,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

add_para(doc3, 'Señores', size=11, bold=True, space_after=2)
add_para(doc3, 'SUBTERRÁNEOS DE BUENOS AIRES S.E.', size=11, bold=True, space_after=2)
add_para(doc3, 'Presente', size=11, space_after=12)

add_para(doc3, 'Ref.: Licitación Privada N° 410/26 – Implementación integral de infraestructura '
         'de redes y telecomunicaciones en edificio Balcarce N° 340', size=11, bold=True, space_after=2)
add_para(doc3, 'DECLARACIÓN JURADA - PLAZO DE GARANTÍA', size=12, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(doc3, 'De nuestra consideración:', size=11, space_after=8)

add_para(doc3, 'Por la presente, en mi carácter de representante legal de SOFTWARE BY DESIGN S.A. '
         '(CUIT 30-70894532-0), declaro bajo juramento que la empresa ofrece para la totalidad '
         'de los trabajos y provisiones correspondientes a la Licitación Privada N° 410/26 '
         'el siguiente plazo de garantía:',
         size=11, space_after=12)

add_para(doc3, 'PLAZO DE GARANTÍA: DOCE (12) MESES', size=13, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para(doc3, 'El plazo de garantía se computará a partir de la fecha del Acta de Recepción '
         'Provisoria de la totalidad de los trabajos, conforme lo establecido en el Pliego '
         'Único de Bases y Condiciones.',
         size=11, space_after=8)

add_para(doc3, 'Durante el período de garantía, SOFTWARE BY DESIGN S.A. se compromete a:',
         size=11, space_after=4)

garantias = [
    'Reparar o reemplazar sin cargo alguno para SBASE cualquier defecto de materiales, '
    'mano de obra o funcionamiento que se manifieste en las instalaciones ejecutadas.',
    'Atender los reclamos de garantía dentro de las cuarenta y ocho (48) horas hábiles '
    'de recibida la notificación correspondiente.',
    'Garantizar el correcto funcionamiento de la totalidad del cableado estructurado, '
    'fibra óptica, equipamiento activo y demás componentes provistos e instalados.',
    'Mantener la certificación de todos los enlaces de cobre y fibra óptica conforme a '
    'las normas aplicables (ANSI/TIA-568.2-D, ISO/IEC 11801).',
]

for g in garantias:
    p = doc3.add_paragraph()
    run = p.add_run('    \u2022  ' + g)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Cm(1)

add_para(doc3, '', space_after=4)
add_para(doc3, 'El plazo de garantía ofrecido cumple con lo establecido en la documentación '
         'del llamado y en la normativa vigente, no siendo inferior a los plazos allí indicados.',
         size=11, space_after=8)

add_para(doc3, 'Sin otro particular, saludo a ustedes muy atentamente.',
         size=11, space_after=24)

add_para(doc3, '', space_after=24)
add_para(doc3, '_________________________________', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc3, 'Hernán Hamra', size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc3, 'Representante Legal', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_para(doc3, 'Software By Design S.A.', size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

out3 = carpeta_b + r'\9 DDJJ Plazo de Garantía.docx'
doc3.save(out3)
print(f'  Guardado: {out3}')

print("\n=== TODO LISTO ===")
