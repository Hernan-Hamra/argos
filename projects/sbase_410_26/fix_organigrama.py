import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from lxml import etree

path = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica\3 Organigrama del Proyecto.docx'

doc = Document(path)
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_para(p):
    elem = p._element
    pPr = elem.find('{%s}pPr' % ns)
    for child in list(elem):
        if child.tag != '{%s}pPr' % ns:
            elem.remove(child)
    if pPr is not None:
        numPr_elem = pPr.find('{%s}numPr' % ns)
        if numPr_elem is not None:
            pPr.remove(numPr_elem)

def set_text(p, text, bold=False, size=11):
    clear_para(p)
    elem = p._element
    new_run = etree.SubElement(elem, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def add_para(doc, after_elem, text, bold=False, size=11):
    new_p = etree.Element('{%s}p' % ns)
    after_elem.addnext(new_p)
    pPr = etree.SubElement(new_p, '{%s}pPr' % ns)
    spacing = etree.SubElement(pPr, '{%s}spacing' % ns)
    spacing.set('{%s}after' % ns, '40')
    spacing.set('{%s}line' % ns, '276')
    spacing.set('{%s}lineRule' % ns, 'auto')
    new_run = etree.SubElement(new_p, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, 'Calibri')
    rFonts.set('{%s}hAnsi' % ns, 'Calibri')
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return new_p

# Print current state for reference
print("=== Estado actual ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'[{i}] {p.text[:100]}')

# --- Fix paragraph 22: "3. Director de Obra" -> bold, keep as is ---
set_text(doc.paragraphs[22], '3. Director de Obra', bold=True, size=11)

# --- Fix paragraph 23: "Leonardo Martinez." -> add full description ---
set_text(doc.paragraphs[23], 'Leonardo Martinez', bold=False, size=11)

# Now I need to add experience text for Leonardo after paragraph 23
# and fix the empty paragraphs 24-25, then fix numbering on 26+

# Check what's at 24, 25
print(f"\n[24] = '{doc.paragraphs[24].text}'")
print(f"[25] = '{doc.paragraphs[25].text}'")

# Paragraph 24: add Leonardo's description
set_text(doc.paragraphs[24],
    'Director de obra con amplia trayectoria en la ejecución de proyectos de '
    'infraestructura de redes y telecomunicaciones. Especialista en cableado '
    'estructurado, fibra óptica, instalación de racks y canalizaciones. '
    'Experiencia comprobada en la dirección de equipos técnicos de instalación '
    'para proyectos de gran envergadura en organismos públicos y empresas privadas, '
    'incluyendo obras de cableado estructurado Cat6/Cat6A, tendido de fibra óptica, '
    'montaje de datacenter, instalación de sistemas de CCTV, control de acceso y '
    'telefonía IP. Responsable de la coordinación operativa en campo, supervisión '
    'de calidad de las instalaciones, cumplimiento de normas y plazos de ejecución.',
    bold=False, size=10)

# Paragraph 25: empty (spacer before equipo)
set_text(doc.paragraphs[25], '', bold=False, size=11)

# --- Fix paragraph 26: "3. Equipo Técnico" -> "4. Equipo Técnico" ---
set_text(doc.paragraphs[26], '4. Equipo Técnico de Instalación', bold=True, size=11)

# --- Fix sub-numbering ---
# [28] 3.1. -> 4.1.
set_text(doc.paragraphs[28], '4.1. Esteban Agustín Marcuzzi - DNI 23.515.129', bold=True, size=11)
# [32] 3.2. -> 4.2.
set_text(doc.paragraphs[32], '4.2. José María Rodriguez - DNI 23.418.511', bold=True, size=11)
# [36] 3.3. -> 4.3.
set_text(doc.paragraphs[36], '4.3. Emanuel Andrés Delpino - DNI 41.432.997', bold=True, size=11)

# --- Update org chart ---
# Chart starts at paragraph 56, update to include Leonardo Martinez
chart_lines = [
    ('              ┌───────────────────────────────────┐', False, 9),
    ('              │       Director de Proyecto         │', False, 9),
    ('              │          Hernán Hamra              │', False, 9),
    ('              └────────────────┬──────────────────┘', False, 9),
    ('                               │', False, 9),
    ('              ┌────────────────┴──────────────────┐', False, 9),
    ('              │        Director Técnico            │', False, 9),
    ('              │        Marcelo Hamra               │', False, 9),
    ('              └────────────────┬──────────────────┘', False, 9),
    ('                               │', False, 9),
    ('              ┌────────────────┴──────────────────┐', False, 9),
    ('              │        Director de Obra            │', False, 9),
    ('              │      Leonardo Martinez             │', False, 9),
    ('              └────────────────┬──────────────────┘', False, 9),
    ('                               │', False, 9),
    ('              ┌────────────────┴──────────────────┐', False, 9),
    ('              │     Equipo Técnico de Instalación  │', False, 9),
    ('              ├───────────────────────────────────┤', False, 9),
    ('              │  Esteban Marcuzzi - Téc. Electr.  │', False, 9),
    ('              │  José M. Rodriguez - Ing. Redes   │', False, 9),
    ('              │  Emanuel Delpino - Téc. Instalador │', False, 9),
    ('              └───────────────────────────────────┘', False, 9),
]

# Old chart goes from paragraph 56 to 72 (17 lines)
# New chart has 22 lines - need to update existing and add extras

# Update existing chart paragraphs (56-72 = 17 paragraphs)
chart_start = 56
for i in range(min(17, len(chart_lines))):
    text, bold, size = chart_lines[i]
    set_text(doc.paragraphs[chart_start + i], text, bold=bold, size=size)

# Add remaining chart lines after paragraph 72
if len(chart_lines) > 17:
    last_elem = doc.paragraphs[72]._element
    for text, bold, size in chart_lines[17:]:
        last_elem = add_para(doc, last_elem, text, bold=bold, size=size)

doc.save(path)
print(f'\nGuardado OK: {path}')
