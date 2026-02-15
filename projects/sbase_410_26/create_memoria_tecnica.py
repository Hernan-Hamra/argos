import sys, os, shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
from copy import deepcopy

template = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\DDJJ\14 Declaración Jurada conocimiento Pliegos.docx'
original_mem = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica\Memoria_Tecnica_LP41026_FINAL ejemplo.docx'
output_temp = r'C:\Users\HERNAN\argos\6 Memoria Técnica y Plan de Trabajo.docx'
output = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica\6 Memoria Técnica y Plan de Trabajo.docx'

# Start from template (has SBD header/footer)
doc = Document(template)

# Remove all existing paragraphs from template body
body = doc._element.body
for p in list(body.findall(qn('w:p'))):
    body.remove(p)

# ---- Helper functions ----
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h

def add_para(doc, text, bold=False, size=10, align=None, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph()
    run = p.add_run('    \u2022  ' + text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Cm(1)
    return p

def add_table_row(table, cells_text, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.bold = bold or header
        if header:
            shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
            shading.set(qn('w:fill'), '1F3A5F')
            shading.set(qn('w:val'), 'clear')
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row

# ==================================================================
# DOCUMENT CONTENT
# ==================================================================

# ---- PORTADA ----
add_para(doc, '', size=10)
add_para(doc, '', size=10)
add_para(doc, 'MEMORIA TÉCNICA', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'PROPUESTA DE SOLUCIÓN, METODOLOGÍA Y PLAN DE EJECUCIÓN', bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '', size=10)
add_para(doc, 'Licitación Privada N° 410/26 - SBASE', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '"Implementación integral de infraestructura de redes y telecomunicaciones', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'en edificio Balcarce N° 340"', size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, '', size=10)
add_para(doc, '', size=10)

# Info box
add_para(doc, 'Empresa Oferente: Software By Design S.A.', size=11)
add_para(doc, 'CUIT: 30-71190824-9', size=11)
add_para(doc, 'Director de Proyecto: Hernán Hamra - Apoderado', size=11)
add_para(doc, 'Director Técnico: Marcelo Ariel Hamra - Presidente', size=11)
add_para(doc, 'Fecha: 18 de febrero de 2026', size=11)
add_para(doc, '', size=10)

# Page break
doc.add_page_break()

# ---- 1. INTRODUCCIÓN Y OBJETO ----
add_heading(doc, '1. Introducción y Objeto', level=1)
add_para(doc,
    'Software By Design S.A. presenta la siguiente Memoria Técnica en el marco de la '
    'Licitación Privada N° 410/26 convocada por Subterráneos de Buenos Aires S.E. (SBASE), '
    'para la "Implementación integral de infraestructura de redes y telecomunicaciones '
    'en edificio Balcarce N° 340".')
add_para(doc,
    'Este documento describe en detalle la solución técnica propuesta, la arquitectura '
    'de red diseñada, el equipamiento seleccionado, la metodología de ejecución, el '
    'plan de trabajo con su cronograma asociado, y los protocolos de certificación y '
    'pruebas que garantizarán la calidad y el correcto funcionamiento de la totalidad '
    'de los sistemas implementados.')
add_para(doc,
    'La propuesta contempla un enfoque integral que abarca la provisión de materiales y '
    'equipamiento, la ejecución de la obra física (canalización y cableado), la instalación '
    'y configuración de los equipos activos de red, telefonía IP y WiFi, la certificación '
    'técnica de cada componente, y la entrega de documentación completa conforme a obra. '
    'El sistema será entregado completamente operativo, listo para su uso inmediato.')

# ---- 2. DESCRIPCIÓN DEL SITIO ----
add_heading(doc, '2. Descripción del Sitio', level=1)
add_para(doc,
    'El edificio Balcarce N° 340, sede administrativa de SBASE, consta de los siguientes '
    'niveles que serán abarcados por la infraestructura de telecomunicaciones:')
add_bullet(doc, 'Subsuelo: Área técnica y de servicios')
add_bullet(doc, 'Planta Baja: Recepción y oficinas operativas')
add_bullet(doc, 'Primer Piso: Oficinas administrativas')
add_bullet(doc, 'Segundo Piso: Oficinas y dirección')
add_bullet(doc, 'Terraza: Equipamiento auxiliar (vinculado al segundo piso)')
add_para(doc,
    'Cada nivel contará con su propio rack de telecomunicaciones equipado con switch de '
    'acceso, patch panel, organizadores, UPS de respaldo y bandeja de fibra óptica para '
    'el enlace troncal. La distribución vertical se resolverá mediante un backbone de '
    'fibra óptica que interconectará todos los pisos, garantizando alta velocidad y '
    'escalabilidad en la comunicación entre sectores.')

# ---- 3. MARCO NORMATIVO ----
add_heading(doc, '3. Marco Normativo y Estándares Aplicables', level=1)
add_para(doc,
    'La totalidad de la obra será ejecutada en cumplimiento de los estándares internacionales '
    'vigentes para infraestructura de telecomunicaciones:')

add_heading(doc, '3.1. Cableado de cobre - Categoría 6A', level=2)
add_bullet(doc, 'ANSI/TIA-568.2-D: Estándar de cableado de telecomunicaciones para edificios comerciales')
add_bullet(doc, 'ISO/IEC 11801-1: Cableado genérico para instalaciones de clientes')
add_bullet(doc, 'ANSI/TIA-569-D: Estándar de canalizaciones y espacios de telecomunicaciones')
add_bullet(doc, 'ANSI/TIA-606-C: Estándar de administración de infraestructura de telecomunicaciones')
add_bullet(doc, 'ANSI/TIA-607-C: Vinculación y referencia de potencial para telecomunicaciones')
add_para(doc,
    'Todos los enlaces de cobre categoría 6A soportarán velocidades de hasta 10 Gbps '
    '(10GBASE-T) con un ancho de banda de 500 MHz, cumpliendo los parámetros eléctricos '
    'de atenuación, NEXT, PS-NEXT, FEXT, PS-ELFEXT y Return Loss exigidos por la norma.')

add_heading(doc, '3.2. Backbone de fibra óptica', level=2)
add_bullet(doc, 'ISO/IEC 11801-1: Especificaciones para cableado de fibra óptica')
add_bullet(doc, 'TIA-568.3-D: Componentes de cableado de fibra óptica')
add_bullet(doc, 'IEEE 802.3: Estándar Ethernet para transmisión óptica')
add_para(doc,
    'El backbone vertical se implementará con fibra óptica multimodo OM3 de 6 hilos, '
    'capaz de soportar 10GBASE-SR hasta 300 metros. Las fusiones se ejecutarán con '
    'empalmadora de arco eléctrico y se verificarán mediante OTDR y medición de potencia óptica.')

add_heading(doc, '3.3. Instalación eléctrica y seguridad', level=2)
add_bullet(doc, 'IEC 62040-1: Requisitos de seguridad para UPS')
add_bullet(doc, 'Reglamentación AEA 90364: Instalaciones eléctricas en inmuebles')

# ---- 4. SOLUCIÓN TÉCNICA PROPUESTA ----
add_heading(doc, '4. Solución Técnica Propuesta', level=1)
add_para(doc,
    'La solución se estructura en capas funcionales, asegurando modularidad, '
    'escalabilidad y facilidad de mantenimiento. A continuación se describe cada '
    'subsistema que compone la infraestructura integral.')

# 4.1 Arquitectura
add_heading(doc, '4.1. Arquitectura General de Red', level=2)
add_para(doc,
    'Se propone una arquitectura jerárquica de dos niveles (core/distribución y acceso), '
    'administrada de forma centralizada mediante el ecosistema UniFi de Ubiquiti. '
    'El gateway central UniFi Cloud Gateway Fiber actuará como núcleo de la red, '
    'proporcionando routing, firewall, IDS/IPS, y gestión unificada de todos los '
    'dispositivos de red (switches, access points).')
add_para(doc,
    'Los switches de acceso UniFi Pro 48 PoE, distribuidos por piso, proveerán '
    'conectividad Gigabit a los puestos de trabajo y alimentación PoE+ a los access '
    'points y teléfonos IP. La interconexión entre switches se realizará mediante '
    'uplinks de 10 Gbps SFP+ sobre el backbone de fibra óptica, eliminando cuellos '
    'de botella en la comunicación entre pisos.')
add_para(doc,
    'La segmentación lógica se implementará mediante VLANs diferenciadas:')
add_bullet(doc, 'VLAN de Datos: Tráfico de estaciones de trabajo y servidores')
add_bullet(doc, 'VLAN de Voz: Tráfico de telefonía IP con priorización QoS')
add_bullet(doc, 'VLAN WiFi Corporativa: Dispositivos inalámbricos autenticados')
add_bullet(doc, 'VLAN WiFi Invitados: Red aislada para visitantes')
add_bullet(doc, 'VLAN de Gestión: Administración de equipos de red')

# 4.2 Infraestructura pasiva
add_heading(doc, '4.2. Infraestructura Pasiva', level=2)
add_para(doc, 'Cableado Estructurado', bold=True)
add_para(doc,
    'Se ejecutará el tendido de cableado estructurado categoría 6A (UTP) '
    'desde cada rack de piso hasta los puestos de trabajo. Cada puesto contará con '
    'un mínimo de dos bocas (datos y telefonía), terminadas en jacks RJ45 Cat6A sobre '
    'face plates identificados. El cableado soportará velocidades de 10 Gbps con un '
    'ancho de banda de 500 MHz.')

add_para(doc, 'Canalización', bold=True)
add_para(doc,
    'La infraestructura de soporte se resolverá mediante bandejas portacables metálicas '
    'con tapa para las distribuciones horizontales, montantes verticales para el backbone '
    'entre pisos, y canaletas o caños de PVC para las bajadas a los puestos de trabajo. '
    'Se respetarán las separaciones reglamentarias respecto de instalaciones eléctricas.')

add_para(doc, 'Racks de Telecomunicaciones', bold=True)
add_para(doc,
    'Se instalarán racks de piso de 42 unidades de rack (42U) en cada nivel del edificio. '
    'Cada rack incluirá:')
add_bullet(doc, 'Patch panels Cat6A de 24 puertos para terminación de cableado horizontal')
add_bullet(doc, 'Organizadores horizontales de cables')
add_bullet(doc, 'Bandeja de empalme de fibra óptica para el backbone vertical')
add_bullet(doc, 'Ventilación (bandeja de ventiladores)')
add_bullet(doc, 'Puertas con cerradura y paneles laterales desmontables')

add_para(doc, 'Backbone de Fibra Óptica', bold=True)
add_para(doc,
    'La interconexión entre racks se realizará mediante cables de fibra óptica multimodo '
    'OM3 de 6 hilos, tendidos por los montantes verticales del edificio. En cada extremo '
    'se terminarán en bandejas de empalme con pigtails LC/UPC fusionados. Esta '
    'infraestructura garantiza un ancho de banda de 10 Gbps entre pisos con posibilidad '
    'de escalabilidad futura a 40 Gbps.')

# 4.3 Equipamiento activo
add_heading(doc, '4.3. Equipamiento Activo de Networking', level=2)

add_para(doc, 'Gateway / Router Central: UniFi Cloud Gateway Fiber (UCG-Fiber)', bold=True)
add_bullet(doc, '2 puertos WAN 10G SFP+ y 1 puerto WAN 10 GbE RJ45')
add_bullet(doc, '4 puertos LAN 2.5 GbE RJ45')
add_bullet(doc, 'Firewall stateful con inspección de capa 7 (DPI)')
add_bullet(doc, 'Sistema de prevención de intrusiones (IDS/IPS) a 5 Gbps')
add_bullet(doc, 'Filtrado de contenido, bloqueo de aplicaciones, bloqueo por regiones')
add_bullet(doc, 'Gestión centralizada UniFi: Network, Protect, Access, Talk')
add_bullet(doc, 'Soporte para más de 50 dispositivos administrados y 500+ usuarios simultáneos')

add_para(doc, 'Switches de Acceso: UniFi Pro 48 PoE (USW-Pro-48-POE)', bold=True)
add_bullet(doc, '48 puertos GbE RJ45 con PoE+ (802.3at) y PoE++ (802.3bt)')
add_bullet(doc, '4 puertos uplink 10G SFP+ para backbone de fibra')
add_bullet(doc, 'Capacidad de switching: 176 Gbps, throughput: 88 Gbps')
add_bullet(doc, 'Presupuesto PoE: 600W por switch')
add_bullet(doc, 'Layer 3: Inter-VLAN routing, DHCP server/relay, routing estático')
add_bullet(doc, 'Soporte de hasta 1,000 VLANs, tabla MAC de 16,000 entradas')
add_bullet(doc, 'Factor de forma: rack 1U, redundancia con fuente DC backup')

# 4.4 WiFi
add_heading(doc, '4.4. Sistema de Conectividad Inalámbrica', level=2)
add_para(doc, 'Access Points: UniFi U6 Long-Range (U6-LR)', bold=True)
add_bullet(doc, 'WiFi 6 (802.11ax) con 8 spatial streams')
add_bullet(doc, 'Banda 5 GHz: 4x4 MU-MIMO, hasta 2.4 Gbps (BW160)')
add_bullet(doc, 'Banda 2.4 GHz: 4x4 SU-MIMO, hasta 600 Mbps (BW40)')
add_bullet(doc, 'Cobertura por AP: 185 m2, capacidad: 350+ clientes simultáneos')
add_bullet(doc, 'Alimentación: PoE+ (802.3at) directamente desde el switch')
add_bullet(doc, 'Protección IP54 para ambientes diversos')
add_bullet(doc, 'Montaje en cielorraso con soporte profesional incluido')
add_para(doc,
    'Los access points serán distribuidos estratégicamente en cada piso para garantizar '
    'cobertura total sin zonas muertas. Se configurarán SSIDs diferenciados para la red '
    'corporativa (con autenticación 802.1X/WPA3) y para invitados (con portal cautivo). '
    'El roaming entre APs será transparente mediante la gestión centralizada del gateway.')

# 4.5 Telefonía IP
add_heading(doc, '4.5. Sistema de Telefonía IP', level=2)
add_para(doc, 'Central IP PBX: Grandstream UCM6300 Audio Series', bold=True)
add_bullet(doc, 'Capacidad: hasta 1,500 usuarios y 200 llamadas simultáneas')
add_bullet(doc, 'Basada en Asterisk 16 (plataforma abierta y extensible)')
add_bullet(doc, 'Interfaces de red: 3x Gigabit RJ45 con PoE+ integrado y NAT router')
add_bullet(doc, 'Codec de alta definición Opus con resiliencia a pérdida de paquetes')
add_bullet(doc, 'Plataforma de conferencias y reuniones integrada')
add_bullet(doc, 'Aprovisionamiento automático (zero-config) de terminales Grandstream')
add_bullet(doc, 'Gestión cloud vía GDMS, API para integración con sistemas de terceros')
add_bullet(doc, 'Seguridad: secure boot, certificado por equipo, encriptación SRTP/TLS')

add_para(doc, 'Teléfonos IP: Grandstream GRP2601P', bold=True)
add_bullet(doc, '2 cuentas SIP, 2 líneas')
add_bullet(doc, 'Conferencia de 5 participantes')
add_bullet(doc, 'Alimentación PoE integrada (sin fuente externa)')
add_bullet(doc, 'Reducción de ruido por IA')
add_bullet(doc, 'Soporte EHS para headsets Plantronics, Jabra, Sennheiser')
add_bullet(doc, 'Aprovisionamiento masivo zero-touch vía GDMS')

add_para(doc, 'Gateway Analógico: Grandstream HT841/HT881', bold=True)
add_bullet(doc, 'HT841: 4 puertos FXO + 1 FXS / HT881: 8 puertos FXO + 1 FXS')
add_bullet(doc, 'Permite interconectar líneas telefónicas analógicas existentes con la central IP')
add_bullet(doc, 'Soporte de fax T.38, lifeline en corte de energía')
add_bullet(doc, '2x Gigabit Ethernet con NAT router integrado')
add_bullet(doc, 'Encriptación AES con certificado por equipo')

add_para(doc,
    'La solución de telefonía se configurará con internos individuales, grupos de '
    'llamada por sector, colas de atención, IVR (menú de opciones), grabación de '
    'llamadas, y reglas de salida diferenciadas. Se priorizará el tráfico de voz '
    'mediante QoS (DSCP/CoS) en toda la red para garantizar calidad de audio sin '
    'cortes ni latencia.')

# 4.6 UPS
add_heading(doc, '4.6. Sistema de Energía Ininterrumpida', level=2)
add_para(doc, 'UPS: Kaise Online Monofásico 1-3 kVA Rack', bold=True)
add_bullet(doc, 'Tecnología: Doble conversión online (zero transfer time)')
add_bullet(doc, 'Factor de potencia de salida: 1.0 (máxima eficiencia)')
add_bullet(doc, 'Rango de entrada: 110-300 VAC (tolerancia extendida)')
add_bullet(doc, 'Salida: Onda senoidal pura, regulación +/-1%')
add_bullet(doc, 'Eficiencia en modo AC: hasta 92%')
add_bullet(doc, 'Comunicaciones: USB, RS232, slot para tarjeta SNMP')
add_bullet(doc, 'Protecciones: EPO, sobrecarga con bypass automático, compatible con generador')
add_bullet(doc, 'Display LCD color rotable')
add_para(doc,
    'Cada rack del edificio contará con un UPS dedicado que asegurará la continuidad '
    'operativa del equipamiento de red ante cortes o variaciones de energía eléctrica. '
    'La tecnología de doble conversión online garantiza un tiempo de transferencia '
    'de 0 ms, protegiendo los equipos de cualquier perturbación eléctrica.')

# ---- 5. DETALLE DE EQUIPAMIENTO ----
add_heading(doc, '5. Detalle de Equipamiento Propuesto', level=1)
add_para(doc, 'A continuación se detalla el equipamiento principal previsto para la obra:')

# Table
table = doc.add_table(rows=1, cols=4)
# table.style = 'Table Grid'  # not available in template
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Header
hdr = table.rows[0]
headers = ['Componente', 'Modelo / Marca', 'Características Principales', 'Cant.']
for i, txt in enumerate(headers):
    cell = hdr.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), '1F3A5F')
    shading.set(qn('w:val'), 'clear')

rows_data = [
    ['Gateway / Router', 'UniFi Cloud Gateway Fiber\n(UCG-Fiber)', '10G SFP+, IDS/IPS 5Gbps, DPI, gestión centralizada UniFi', '1'],
    ['Switches de Acceso', 'UniFi Pro 48 PoE\n(USW-Pro-48-POE)', '48p GbE PoE+ 600W, 4x10G SFP+, L3, 176 Gbps, 1U rack', '10'],
    ['Access Points WiFi 6', 'UniFi U6 Long-Range\n(U6-LR)', '802.11ax, 2.4Gbps 5GHz, 185m2 cobertura, PoE+, IP54', '17'],
    ['Central IP PBX', 'Grandstream UCM6300A', 'Hasta 1500 usuarios, 200 llamadas simult., Asterisk 16, SRTP', '1'],
    ['Teléfonos IP', 'Grandstream GRP2601P', '2 SIP, PoE integrado, conferencia 5 partic., zero-touch', '100'],
    ['Gateway FXO', 'Grandstream HT841', '4 FXO + 1 FXS, 2x GbE, T.38 fax, lifeline', 'c/n'],
    ['UPS Online', 'Kaise 1-3kVA Rack', 'Doble conversión, 0ms transfer, PF 1.0, LCD, SNMP', '5'],
    ['Rack 42U', 'Rack de piso estándar', '42U, puertas con cerradura, bandejas, ventilación', '5'],
    ['Patch Panel Cat6A', '24 puertos Cat6A', 'Terminación Krone/110, montaje rack 19"', '30'],
    ['Cable UTP Cat6A', 'Bobina 305m', 'UTP, 500 MHz, 10GBASE-T, LSZH', 'c/n'],
    ['Fibra óptica OM3', '6 hilos multimodo', '10GBASE-SR, interior, LSZH', 'c/n'],
    ['Patch cords Cat6A', '0.5m y 1m', 'Cat6A certificados, distintas longitudes', '230'],
]

for row_data in rows_data:
    row = table.add_row()
    for i, txt in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(8)
        run.font.name = 'Calibri'

add_para(doc, '')
add_para(doc, 'Nota: Las fichas técnicas completas (datasheets) de cada equipo se adjuntan como documentación complementaria.', size=9)

# ---- 6. METODOLOGÍA DE EJECUCIÓN ----
add_heading(doc, '6. Metodología de Ejecución', level=1)
add_para(doc,
    'La ejecución del proyecto se organizará en fases secuenciales con solapamiento '
    'controlado, optimizando los tiempos sin comprometer la calidad. Cada fase contará '
    'con puntos de control y verificación antes de avanzar a la siguiente.')

add_heading(doc, '6.1. Fase 1 - Ingeniería de Detalle y Relevamiento (Días 1 a 5)', level=2)
add_bullet(doc, 'Relevamiento in situ del edificio: recorrido de pisos, verificación de canalizaciones existentes, ubicación de racks')
add_bullet(doc, 'Definición de recorridos de bandejas y montantes')
add_bullet(doc, 'Planos de detalle con ubicación de cada boca de red, AP y rack')
add_bullet(doc, 'Esquema de direccionamiento IP, VLANs y numeración de internos')
add_bullet(doc, 'Coordinación con el área de mantenimiento de SBASE para accesos y horarios de obra')

add_heading(doc, '6.2. Fase 2 - Aprovisionamiento de Materiales (Días 1 a 7)', level=2)
add_para(doc,
    'Esta fase se ejecuta en paralelo con la ingeniería de detalle. Comprende la '
    'adquisición y logística de todo el equipamiento y materiales, incluyendo verificación '
    'de stock, importaciones si fueran necesarias, y recepción en depósito con control '
    'de calidad e inventario.')

add_heading(doc, '6.3. Fase 3 - Obra Civil: Canalización y Montaje de Racks (Días 6 a 15)', level=2)
add_bullet(doc, 'Instalación de bandejas portacables metálicas con tapa en distribuciones horizontales')
add_bullet(doc, 'Ejecución de montantes verticales entre pisos')
add_bullet(doc, 'Instalación de canaletas y caños para bajadas a puestos de trabajo')
add_bullet(doc, 'Montaje de racks de 42U con todos sus accesorios')
add_bullet(doc, 'Tendido de alimentación eléctrica dedicada a cada rack con protección termomagnética y diferencial')
add_bullet(doc, 'Verificación de condiciones eléctricas existentes para los racks')

add_heading(doc, '6.4. Fase 4 - Cableado Estructurado (Días 12 a 25)', level=2)
add_bullet(doc, 'Tendido de cables UTP Cat6A desde cada rack hacia los puestos de trabajo')
add_bullet(doc, 'Terminación de jacks RJ45 Cat6A en face plates en puestos de usuario')
add_bullet(doc, 'Terminación en patch panels Cat6A de 24 puertos en racks')
add_bullet(doc, 'Tendido de backbone de fibra óptica OM3 (6 hilos) por montantes verticales')
add_bullet(doc, 'Fusión de fibras con empalmadora de arco eléctrico en bandejas ópticas')
add_bullet(doc, 'Etiquetado provisional de todos los puntos conforme norma TIA-606')

add_heading(doc, '6.5. Fase 5 - Instalación de Equipamiento Activo (Días 22 a 30)', level=2)
add_bullet(doc, 'Instalación del gateway UniFi Cloud Gateway Fiber en rack principal')
add_bullet(doc, 'Instalación de switches UniFi Pro 48 PoE en cada rack de piso')
add_bullet(doc, 'Montaje de access points UniFi U6-LR en cielorrasos')
add_bullet(doc, 'Instalación de central telefónica Grandstream UCM6300A')
add_bullet(doc, 'Conexión de gateway FXO Grandstream HT841 a líneas analógicas existentes')
add_bullet(doc, 'Distribución e instalación de teléfonos IP GRP2601P en cada puesto')
add_bullet(doc, 'Instalación de UPS Kaise en cada rack')
add_bullet(doc, 'Conexión de patch cords en patch panels y equipos')

add_heading(doc, '6.6. Fase 6 - Configuración y Puesta en Marcha (Días 28 a 40)', level=2)
add_bullet(doc, 'Configuración del gateway: WAN, firewall, VLANs, DHCP, DNS, IDS/IPS, DPI')
add_bullet(doc, 'Configuración de switches: VLANs, trunk ports, access ports, PoE, STP, QoS')
add_bullet(doc, 'Configuración de access points: SSIDs, autenticación, roaming, bandas')
add_bullet(doc, 'Programación de la central telefónica: internos, grupos, IVR, colas, grabación')
add_bullet(doc, 'Configuración de QoS end-to-end para priorización de voz')
add_bullet(doc, 'Aprovisionamiento automático de teléfonos IP vía GDMS')
add_bullet(doc, 'Verificación de backup de configuraciones de todos los equipos')

add_heading(doc, '6.7. Fase 7 - Certificación, Pruebas y Documentación (Días 38 a 45)', level=2)
add_bullet(doc, 'Certificación de cada enlace de cobre Cat6A con certificador Fluke o equivalente')
add_bullet(doc, 'Verificación de enlaces de fibra óptica mediante OTDR y medición de potencia')
add_bullet(doc, 'Pruebas de conectividad IP extremo a extremo en todos los puntos')
add_bullet(doc, 'Pruebas de telefonía: llamadas internas, externas, conferencias, IVR')
add_bullet(doc, 'Pruebas de cobertura WiFi: site survey post-instalación, roaming, throughput')
add_bullet(doc, 'Verificación de funcionamiento de UPS: test de autonomía y transferencia')
add_bullet(doc, 'Elaboración de documentación conforme a obra')
add_bullet(doc, 'Entrega formal con acta de recepción provisoria')

# ---- 7. CRONOGRAMA ----
add_heading(doc, '7. Plan de Trabajo y Cronograma', level=1)
add_para(doc,
    'El plazo total de ejecución será de cuarenta y cinco (45) días corridos contados '
    'a partir de la fecha de la Orden de Inicio. Las fases se ejecutarán con solapamiento '
    'controlado para optimizar el plazo global:')
add_para(doc, '')

# Cronograma table
ctable = doc.add_table(rows=1, cols=3)
ctable.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(['Fase', 'Descripción', 'Plazo']):
    cell = ctable.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(txt)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), '1F3A5F')
    shading.set(qn('w:val'), 'clear')

cron_data = [
    ['1', 'Ingeniería de detalle y relevamiento', 'Día 1 a 5'],
    ['2', 'Aprovisionamiento de materiales', 'Día 1 a 7'],
    ['3', 'Obra civil: canalización y montaje de racks', 'Día 6 a 15'],
    ['4', 'Cableado estructurado y fibra óptica', 'Día 12 a 25'],
    ['5', 'Instalación de equipamiento activo', 'Día 22 a 30'],
    ['6', 'Configuración y puesta en marcha', 'Día 28 a 40'],
    ['7', 'Certificación, pruebas y documentación', 'Día 38 a 45'],
]
for row_data in cron_data:
    row = ctable.add_row()
    for i, txt in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

add_para(doc, '')
add_para(doc, 'Diagrama de Gantt:', bold=True)

# ---- INSERT NEW GANTT IMAGE ----
gantt_img = r'C:\Users\HERNAN\argos\gantt_sbd.png'
if os.path.exists(gantt_img):
    doc.add_picture(gantt_img, width=Inches(6.2))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print('New Gantt image inserted OK')
else:
    add_para(doc, '[Diagrama de Gantt]')

# ---- 8. IDENTIFICACIÓN Y ROTULADO ----
add_heading(doc, '8. Sistema de Identificación y Rotulado', level=1)
add_para(doc,
    'La identificación del cableado se ejecutará conforme a la norma ANSI/TIA-606-C. '
    'Cada punto de la red será identificado de forma única e inequívoca en ambos extremos '
    '(boca de usuario y patch panel) mediante etiquetas permanentes impresas con '
    'rotuladora profesional.')
add_para(doc, 'Esquema de nomenclatura:', bold=True)
add_para(doc, 'El identificador de cada punto tendrá el formato: [PISO]-[RACK]-[BOCA]')
add_para(doc, 'Ejemplo: PB-R1-024 = Planta Baja, Rack 1, Boca 024')
add_para(doc, 'Elementos que serán rotulados:')
add_bullet(doc, 'Tomas de usuario (face plates)')
add_bullet(doc, 'Puertos de patch panels')
add_bullet(doc, 'Patch cords de cobre')
add_bullet(doc, 'Hilos de fibra óptica y bandejas de empalme')
add_bullet(doc, 'Puertos de switches (descripción en configuración)')
add_bullet(doc, 'Racks y bandejas portacables')

# ---- 9. CERTIFICACIÓN ----
add_heading(doc, '9. Protocolo de Certificación y Pruebas', level=1)
add_para(doc,
    'Al finalizar la instalación, se ejecutará un protocolo completo de certificación '
    'y pruebas para verificar el correcto funcionamiento de cada componente del sistema:')

add_heading(doc, '9.1. Certificación de cableado de cobre', level=2)
add_para(doc,
    'Cada enlace de cobre Cat6A será certificado con equipo certificador nivel IIIe '
    '(tipo Fluke DSX o equivalente), verificando los parámetros:')
add_bullet(doc, 'Wire Map (continuidad y polaridad)')
add_bullet(doc, 'Longitud del enlace')
add_bullet(doc, 'Insertion Loss (atenuación)')
add_bullet(doc, 'NEXT y PS-NEXT (paradiafonia)')
add_bullet(doc, 'ACR-F y PS-ACR-F (relación atenuación/paradiafonia)')
add_bullet(doc, 'Return Loss')
add_bullet(doc, 'Propagation Delay y Delay Skew')
add_para(doc, 'Se entregarán los reportes de certificación de cada enlace en formato digital.')

add_heading(doc, '9.2. Verificación de fibra óptica', level=2)
add_bullet(doc, 'Medición de potencia óptica en cada hilo (dBm)')
add_bullet(doc, 'Medición OTDR para verificar fusiones, conectores y pérdidas por tramo')

add_heading(doc, '9.3. Pruebas funcionales', level=2)
add_bullet(doc, 'Conectividad IP punto a punto en todos los puestos')
add_bullet(doc, 'Velocidad de enlace verificada a 1 Gbps en cada boca')
add_bullet(doc, 'Pruebas de telefonía: llamadas internas, externas, conferencia, IVR')
add_bullet(doc, 'Site survey WiFi post-instalación: cobertura, intensidad de señal, roaming')
add_bullet(doc, 'Test de autonomía de UPS y verificación de transferencia')

# ---- 10. DOCUMENTACIÓN ----
add_heading(doc, '10. Documentación Final de Obra', level=1)
add_para(doc, 'Al finalizar la obra se entregará la siguiente documentación:')
add_bullet(doc, 'Planos conforme a obra (as-built) con ubicación de racks, bocas, APs y recorridos de cableado')
add_bullet(doc, 'Diagrama lógico de red con VLANs, direccionamiento IP y topología')
add_bullet(doc, 'Planilla de detalle de bocas: identificación, rack, patch panel, puerto, estado')
add_bullet(doc, 'Reportes de certificación de cableado de cobre (Fluke o equivalente)')
add_bullet(doc, 'Reportes de medición de fibra óptica (OTDR y potencia)')
add_bullet(doc, 'Respaldo de configuración de todos los equipos activos (gateway, switches, PBX)')
add_bullet(doc, 'Manual básico de operación y administración del sistema')
add_bullet(doc, 'Documentación de garantía de los equipos')

# ---- 11. GARANTÍA ----
add_heading(doc, '11. Garantía y Soporte Técnico Post-Entrega', level=1)
add_para(doc,
    'Software By Design S.A. ofrecerá un plazo de garantía de doce (12) meses '
    'contados a partir de la fecha de Recepción Provisoria de la obra, conforme lo '
    'establecido en el Pliego de Especificaciones Técnicas.')
add_para(doc, 'La garantía cubrirá:')
add_bullet(doc, 'Funcionamiento correcto de toda la infraestructura instalada')
add_bullet(doc, 'Reemplazo de materiales o equipos con defectos de fábrica')
add_bullet(doc, 'Soporte técnico para resolución de incidencias')
add_bullet(doc, 'Reparación o reemplazo de puntos de cableado que presenten fallas')
add_para(doc,
    'El soporte técnico se canalizará a través de los contactos del Director de Proyecto '
    '(Hernán Hamra) con respuesta dentro de las 48 horas hábiles de reportado el incidente.')

# ---- 12. EQUIPO DE PROYECTO ----
add_heading(doc, '12. Equipo de Proyecto', level=1)
add_para(doc, 'El proyecto será ejecutado por el siguiente equipo profesional:')

# Team table
ttable = doc.add_table(rows=1, cols=3)
ttable.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(['Rol', 'Nombre', 'Responsabilidad']):
    cell = ttable.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(txt)
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shading = etree.SubElement(cell._element.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), '1F3A5F')
    shading.set(qn('w:val'), 'clear')

team_data = [
    ['Director de Proyecto', 'Hernán Hamra', 'Dirección general, coordinación con SBASE, gestión contractual'],
    ['Director Técnico', 'Marcelo Ariel Hamra', 'Supervisión técnica, arquitectura de red, control de calidad'],
    ['Director de Obra', 'Leonardo Martinez', 'Ejecución en campo, coordinación de equipo técnico'],
    ['Técnico Electrónico', 'Esteban A. Marcuzzi', 'Instalación de redes, cableado, montaje de equipos'],
    ['Ing. Redes y Sistemas', 'José M. Rodriguez', 'Cableado estructurado, fibra óptica, configuración de red'],
    ['Técnico Instalador', 'Emanuel A. Delpino', 'Instalación de infraestructura pasiva y equipamiento'],
]
for row_data in team_data:
    row = ttable.add_row()
    for i, txt in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'

# ---- CIERRE Y FIRMAS ----
add_para(doc, '')
add_para(doc, '')
add_para(doc,
    'La presente Memoria Técnica forma parte de la propuesta de Software By Design S.A. '
    'para la Licitación Privada N° 410/26 de SBASE y refleja el compromiso de la empresa '
    'con la excelencia técnica en la ejecución de proyectos de infraestructura de '
    'telecomunicaciones.', size=10)

add_para(doc, '')
add_para(doc, '')
add_para(doc, '')

# Firmas side by side - use a table
ftable = doc.add_table(rows=4, cols=2)
ftable.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in ftable.rows:
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

ftable.cell(0, 0).paragraphs[0].add_run('_____________________________').font.size = Pt(10)
ftable.cell(0, 1).paragraphs[0].add_run('_____________________________').font.size = Pt(10)
ftable.cell(1, 0).paragraphs[0].add_run('Hernán Hamra').font.size = Pt(10)
ftable.cell(1, 1).paragraphs[0].add_run('Marcelo Ariel Hamra').font.size = Pt(10)
ftable.cell(2, 0).paragraphs[0].add_run('Director de Proyecto').font.size = Pt(9)
ftable.cell(2, 1).paragraphs[0].add_run('Director Técnico').font.size = Pt(9)
ftable.cell(3, 0).paragraphs[0].add_run('Apoderado - SBD S.A.').font.size = Pt(9)
ftable.cell(3, 1).paragraphs[0].add_run('Presidente - SBD S.A.').font.size = Pt(9)

# Remove borders from signature table
for row in ftable.rows:
    for cell in row.cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
        for border in ['top', 'left', 'bottom', 'right']:
            b = etree.SubElement(tcBorders, qn(f'w:{border}'))
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')

doc.save(output_temp)
print(f'Guardado temp: {output_temp}')
import shutil
try:
    shutil.copy2(output_temp, output)
    print(f'Copiado a: {output}')
except Exception as e:
    print(f'No se pudo copiar a destino (cerrar el archivo): {e}')
    print(f'Archivo disponible en: {output_temp}')
print('OK')
