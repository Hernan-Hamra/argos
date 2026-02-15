import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

path = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR\CARPETA B Documentación Técnica\3 Organigrama del Proyecto.docx'

doc = Document(path)
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_para(p):
    elem = p._element
    pPr = elem.find('{%s}pPr' % ns)
    for child in list(elem):
        if child.tag != '{%s}pPr' % ns:
            elem.remove(child)
    if pPr is not None:
        numPr_elem = pPr.find('{%s}numPr' % ns)
        if numPr_elem is not None:
            pPr.remove(numPr_elem)

def set_text(p, text, bold=False, size=11, font='Calibri'):
    clear_para(p)
    elem = p._element
    new_run = etree.SubElement(elem, '{%s}r' % ns)
    rPr = etree.SubElement(new_run, '{%s}rPr' % ns)
    rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
    rFonts.set('{%s}ascii' % ns, font)
    rFonts.set('{%s}hAnsi' % ns, font)
    rFonts.set('{%s}cs' % ns, font)
    sz = etree.SubElement(rPr, '{%s}sz' % ns)
    sz.set('{%s}val' % ns, str(size * 2))
    szCs = etree.SubElement(rPr, '{%s}szCs' % ns)
    szCs.set('{%s}val' % ns, str(size * 2))
    if bold:
        b_elem = etree.SubElement(rPr, '{%s}b' % ns)
    lang = etree.SubElement(rPr, '{%s}lang' % ns)
    lang.set('{%s}val' % ns, 'es-AR')
    t = etree.SubElement(new_run, '{%s}t' % ns)
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

# Find chart paragraphs - they contain box-drawing characters
# Rewrite all chart lines with Consolas font and proper alignment
# Using simpler, well-aligned boxes

chart_lines = [
    '┌─────────────────────────────────┐',
    '│     Director de Proyecto        │',
    '│        Hernán Hamra             │',
    '└───────────────┬─────────────────┘',
    '                │                  ',
    '┌───────────────┴─────────────────┐',
    '│      Director Técnico           │',
    '│      Marcelo Hamra              │',
    '└───────────────┬─────────────────┘',
    '                │                  ',
    '┌───────────────┴─────────────────┐',
    '│      Director de Obra           │',
    '│    Leonardo Martinez            │',
    '└───────────────┬─────────────────┘',
    '                │                  ',
    '┌───────────────┴─────────────────┐',
    '│  Equipo Técnico de Instalación  │',
    '├─────────────────────────────────┤',
    '│ Esteban Marcuzzi - Téc. Electr. │',
    '│ José M. Rodriguez - Ing. Redes  │',
    '│ Emanuel Delpino - Téc. Instal.  │',
    '└─────────────────────────────────┘',
]

# Verify all lines are same width
max_len = max(len(l) for l in chart_lines)
print(f"Max line length: {max_len}")
for i, line in enumerate(chart_lines):
    print(f"  [{i}] len={len(line)} | {line}")

# Find chart start - look for first box drawing char
chart_start = None
chart_end = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt and ('┌' in txt or '│' in txt or '└' in txt or '├' in txt):
        if chart_start is None:
            chart_start = i
        chart_end = i

print(f"\nChart paragraphs: {chart_start} to {chart_end}")
print(f"Existing chart lines: {chart_end - chart_start + 1}")
print(f"New chart lines: {len(chart_lines)}")

# Update existing chart paragraphs with Consolas font
existing_count = chart_end - chart_start + 1

for i in range(len(chart_lines)):
    p_idx = chart_start + i
    if i < existing_count:
        # Update existing paragraph
        set_text(doc.paragraphs[p_idx], chart_lines[i], font='Consolas', size=9)
    else:
        # Need to add new paragraph - shouldn't happen since new has 22 = old has 22
        print(f"WARNING: Need to add extra paragraph for line {i}")

# If old had more lines than new, clear extras
if existing_count > len(chart_lines):
    for i in range(len(chart_lines), existing_count):
        p_idx = chart_start + i
        set_text(doc.paragraphs[p_idx], '', font='Consolas', size=9)

doc.save(path)
print(f'\nGuardado OK')
