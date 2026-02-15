import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

base = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\1 LICITACIONES EN PRESENTACIÓN\3 SBASE CABLEADO ESTRUCTURADO\LICI SBASE CABLEADO 2026\DOC A PRESENTAR'
template = base + r'\DDJJ\14 Declaración Jurada conocimiento Pliegos.docx'
ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag != '{%s}sectPr' % ns:
            body.remove(child)

def add_para(doc, text, bold=False, size=11, align=None, space_after=6,
             space_before=2, font='Calibri', color=None):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = font
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        r_elem = run._element
        rPr = r_elem.find('{%s}rPr' % ns)
        if rPr is None:
            rPr = etree.SubElement(r_elem, '{%s}rPr' % ns)
        rFonts = rPr.find('{%s}rFonts' % ns)
        if rFonts is None:
            rFonts = etree.SubElement(rPr, '{%s}rFonts' % ns)
        rFonts.set('{%s}ascii' % ns, font)
        rFonts.set('{%s}hAnsi' % ns, font)
        rFonts.set('{%s}cs' % ns, font)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run('_' * 60)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(27, 79, 114)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(12)

caratulas = [
    {
        'nombre': 'CARPETA A1',
        'subtitulo': 'Documentación Legal',
        'destino': 'CARPETA A1 Documentación Legal',
        'contenido': [
            '1. Carta de Presentación (Anexo I)',
            '2. Garantía de Oferta',
            '3. DDJJ de Intereses (Anexo II)',
            '4. Constancia de registro y obtención del pliego',
            '5. Acta constitutiva y estatuto',
            '6. Constancia de visita obligatoria',
            '7. Domicilio legal constituido en CABA',
            '8. DDJJ Aptitud para Contratar (Anexo III)',
            '9. DDJJ Impedimentos para Participar (Anexo IV)',
            '10. DDJJ Litigio judicial',
            '11. DDJJ Juicios pendientes',
            '12. Poder del representante',
            '13. Inscripción RIUPP',
            '14. DDJJ Conocimiento y aceptación de Pliegos',
            '15. DDJJ Versiones digitales copia fiel',
            '16. Certificados (deudor alimentario / antecedentes penales)',
        ],
    },
    {
        'nombre': 'CARPETA A2',
        'subtitulo': 'Información Económico-Financiera',
        'destino': 'CARPETA A2 Información Económico-Financiera',
        'contenido': [
            '1. Constancia de inscripción Ingresos Brutos',
            '2. DDJJ IIBB últimos 12 meses con presentación y pago',
            '3. DDJJ IVA últimos 12 meses con constancia de pago',
            '4. Constancia de inscripción AFIP (ARCA)',
            '5. Memoria y Balance últimos 2 ejercicios',
            '6. Constancia presentación balances en IGJ',
            '7. Certificación contable de ventas',
            '8. Listado de bancos y autorización de referencias',
        ],
    },
    {
        'nombre': 'CARPETA B',
        'subtitulo': 'Documentación Técnica',
        'destino': 'CARPETA B Documentación Técnica',
        'contenido': [
            '1. Datos de la empresa',
            '2. Antecedentes en obras similares (DDJJ + Listado + OC)',
            '3. Organigrama y plantel profesional',
            '4. Representante Técnico (CV + Designación + Aceptación)',
            '5. CV Director de Proyecto',
            '6. Memoria Técnica y Plan de Trabajo',
            '7. Lista de proveedores propuestos',
            '8. Nota sobre subcontratistas',
            '9. DDJJ Plazo de Garantía',
            '10. Datasheets y folletos técnicos',
        ],
    },
    {
        'nombre': 'CARPETA C',
        'subtitulo': 'Oferta Económica',
        'destino': 'CARPETA C Oferta Económica',
        'contenido': [
            '1. Fórmula de la Oferta (Anexo V)',
            '2. Planillas de Cotización y Desglose (Anexo VI)',
            '3. Análisis de Precios (Anexo VII)',
        ],
    },
]

for car in caratulas:
    doc = Document(template)
    clear_body(doc)

    add_para(doc, '', space_after=30)
    add_para(doc, '', space_after=30)

    add_para(doc, 'LICITACIÓN PRIVADA N° 410/26', size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
             color=(27, 79, 114))

    add_para(doc, 'SUBTERRÁNEOS DE BUENOS AIRES S.E.', size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
             color=(27, 79, 114))

    add_para(doc, 'Implementación integral de infraestructura de redes', size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, 'y telecomunicaciones en edificio Balcarce N° 340', size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_separator(doc)

    add_para(doc, car['nombre'], size=28, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
             color=(27, 79, 114))

    add_para(doc, car['subtitulo'], size=18, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6,
             color=(93, 109, 126))

    add_separator(doc)

    add_para(doc, 'CONTENIDO:', size=11, bold=True, space_after=8,
             color=(27, 79, 114))

    for item in car['contenido']:
        p = add_para(doc, item, size=10, space_after=3, space_before=1)
        p.paragraph_format.left_indent = Cm(1)

    add_para(doc, '', space_after=30)
    add_separator(doc)
    add_para(doc, 'SOFTWARE BY DESIGN S.A.', size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
             color=(27, 79, 114))
    add_para(doc, 'CUIT 30-70894532-0', size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
             color=(93, 109, 126))
    add_para(doc, 'Febrero 2026', size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
             color=(93, 109, 126))

    out = base + '\\' + car['destino'] + '\\0 Carátula.docx'
    doc.save(out)
    print(f'OK: {car["destino"]}\\0 Carátula.docx')

print('\nTodos guardados.')
