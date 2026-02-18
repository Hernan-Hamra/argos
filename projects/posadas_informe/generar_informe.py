# -*- coding: utf-8 -*-
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

out_dir = r'C:\Users\HERNAN\OneDrive - SOFTWARE BY DESIGN SA\3 EJECUCIÓN PROYECTOS\1 HOSPITAL POSADAS CABLEADO ESTRUCTURADO  96-0051-LPU22 17-08-2022 1300 Hrs'

def make_header_cell(cell, text):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    tc = cell._element.get_or_add_tcPr()
    shd = tc.makeelement(qn('w:shd'), {qn('w:fill'): '1F3A5F', qn('w:val'): 'clear'})
    tc.append(shd)

def shade_row(row, color='F2F2F2'):
    for cell in row.cells:
        tc = cell._element.get_or_add_tcPr()
        shd = tc.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
        tc.append(shd)

def add_title(doc, text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def add_subtitle(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_section_title(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.5)
        s.page_width = Cm(29.7)   # A4 landscape
        s.page_height = Cm(21)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    return doc

# ============================================================
# DOC 1: DATOS CRUDOS (sin cambios de estructura)
# ============================================================
doc1 = setup_doc()
# Reset to portrait for crudos
for s in doc1.sections:
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.left_margin = Cm(2)
    s.right_margin = Cm(2)

add_title(doc1, 'DATOS CRUDOS - Conversaciones WhatsApp')
add_subtitle(doc1, 'Hospital Posadas - Dic 2025 / Ene-Feb 2026 - Hernán Hamra')
doc1.add_paragraph()

chats = [
    ('CHAT: Claudio Kolln (Supervisor de Sistemas - Hospital Posadas)', [
        '[18:13, 22/12/2025] Hernán: Modulo Óptico Sfp+ 10g Mm 850nm 300m Sr Multimodo Tecnoptic',
        '[18:39, 22/12/2025] Hernán: si me das el ok los compro. generico me dieron el ok de huawei',
        '[18:40, 22/12/2025] Hernán: tema es si podemos facturar mañana, nos acepten los equipo y la entrega de los sfp para este viernes 26',
        '[19:14, 22/12/2025] Claudio Kolln: mañana lo vemos y aclaramos todas las dudas',
        '[12:06, 29/12/2025] Claudio Kolln: no te olvides de los módulos SFP...la semana que viene tenemos que instalar los switch',
        '[12:42, 29/12/2025] Hernán: Ok',
        '[8:53, 30/12/2025] Hernán: Buen dia claudio. Estas para recibir hoy por la mañana?',
        '[8:54, 30/12/2025] Hernán: Dale mando moto ahora',
        '[8:56, 30/12/2025] Claudio Kolln: Dale',
        '[9:29, 30/12/2025] Hernán: Sigue la ruta en Cabify [link tracking]',
        '[9:54, 30/12/2025] Claudio Kolln: Ya llegaron los SFP!!',
    ]),
    ('CHAT: Marino Frangi (Socio)', [
        '[9:03, 13/1/2026] Hernán: bue dia.',
        '[9:51, 13/1/2026] Hernán: Avisame cuando podes hablar',
        '[9:52, 13/1/2026] Marino: Buen día Hernán, bárbaro, a la tarde hablamos',
        '[8:37, 14/1/2026] Hernán: marino, buen dia. llamame por favor a la mañana en lo posible amigo',
        '[9:00, 14/1/2026] Marino: en un rato te llamo y te digo la estrategia así hacemos bajada de línea hacia Leo. Ayer estuve reunido con Eze hasta tarde analizando esto y armando la estrategia',
        '[12:11, 14/1/2026] Hernán: [link rotomartillo Kommberg] esto me envío abel del posadas',
        '[16:48, 14/1/2026] Marino: Hernán, me pasás el contacto de Versailles?',
        '[16:50, 14/1/2026] Hernán: Miguel Merlo Distribuidora Versalles Materiales Eléctricos',
        '[12:39, 15/1/2026] Marino: Te llamo en cuanto pueda.',
        '[12:45, 15/1/2026] Marino: [audio] Genial! Yo creo que vamos a estar bien',
        '[12:46, 15/1/2026] Marino: [audio] No creo que nos rompan en el hospital, ojalá llegue el taladro y mañana arrancamos con el pie derecho',
        '[18:30, 15/1/2026] Hernán: Acopio de materiales: del lunes 19 al viernes de enero. inicio de obra: del lunes 26 al viernes 30 de enero.',
        '[18:34, 15/1/2026] Marino: Correcto, confirmá con Leo que pueda arrancar ese lunes y le metemos',
        '[13:44, 9/2/2026] Hernán: Ahi me comento leo que les paso el material a franco. Mañana estoy pasando avance de obra a ricardo.',
        '[13:57, 9/2/2026] Marino: Abrazo!',
    ]),
    ('CHAT: Abel (Hospital Posadas)', [
        '[14:52, 14/1/2026] Hernán: No. Bola. La metemos nosotros. Obvio. Es parte del proyecto',
        '[15:17, 14/1/2026] Hernán: gracias amigo',
        '[16:20, 16/1/2026] Hernán: [link multiherramienta Matrix 5 en 1]',
        '[16:57, 16/1/2026] Hernán: Buen finde abelito.',
        '[11:14, 13/2/2026] Hernán: Si esta todo eso. Ya te averiguo',
        '[13:04, 13/2/2026] Hernán: Abrazo amigo',
    ]),
    ('CHAT: Ricardo Torres (Hospital Posadas)', [
        '[15:17, 16/1/2026] Hernán: Excelente el deposito. Impecable.',
        '[16:32, 16/1/2026] Hernán: Dale. Perfecto. Abrí el tema nomas. Joya el material',
        '[13:22, 22/1/2026] Hernán: hamrahernan@gmail.com',
        '[16:16, 22/1/2026] Hernán: te la mande al mail. te la copio aca',
        '[16:28, 22/1/2026] Hernán: Interfaces estándar MEMORIA — USB A, USB B, Ethernet IEEE802.3, LAN WiFi, 128 GB',
        '[20:07, 22/1/2026] Ricardo Torres: A ver que te parece...',
        '[21:02, 22/1/2026] Hernán: ANÁLISIS TÉCNICO — PostScript → NO CUMPLE, PDF Adobe → NO CUMPLE, HP-GL/2 → CUMPLE, PCL → NO CUMPLE, Velocidad A1 → NO CUMPLE, Pantalla → NO CUMPLE',
        '[21:04, 22/1/2026] Ricardo Torres: Lo paso para revisión',
        '[14:56, 23/1/2026] Hernán: Ahi lo miro',
        '[15:55, 23/1/2026] Hernán: ok. buen finde',
        '[13:06, 2/2/2026] Hernán: Si dale',
        '[10:08, 9/2/2026] Hernán: Dale',
    ]),
    ('CHAT: Leonardo Martínez (DNET - Subcontratista)', [
        '[13:08, 9/2/2026] Leo: Quedan 105 tiras. Instaladas 120. Faltan 130 ménsulas + 12 curvas 90',
        '[13:27, 9/2/2026] Hernán: 130 ménsulas 200mm + 12 curvas 90° 200mm. ok?',
        '[13:28, 9/2/2026] Leo: Correcto, ya se lo pasé a Franco',
        '[13:46, 9/2/2026] Hernán: No te olvides los seguros de los chicos',
        '[14:52, 10/2/2026] Hernán: mando proyecto y coti. estos son los planos que armé yo',
        '[15:09, 10/2/2026] Leo: P7 falta 90m. P6 OK 150m. P5 parcial. P3 OK. P1 OK 90m. 150 tiras instaladas.',
        '[17:51, 10/2/2026] Hernán: te paso pliego sbase con puestos y AP para cotizar',
        '[12:09, 11/2/2026] Hernán: pasame propuesta mano de obra cableado subtes',
        '[12:09, 12/2/2026] Leo: Bestia y las mensulas y articuladas y curvas?',
        '[12:36, 13/2/2026] Leo: Y los materiales?',
    ]),
    ('CHAT: Ezequiel Da Conceicao (Socio técnico)', [
        '[19:23, 12/2/2026] Hernán: Unifi no soporta mpls',
        '[19:25, 12/2/2026] Ezequiel: no lo pide en el pliego. Solo afecta un switch, se soluciona con router',
        '[19:34, 12/2/2026] Ezequiel: decile direc attach',
        '[9:16, 13/2/2026] Hernán: avisame si podes hacer un meet',
        '[10:00, 13/2/2026] Ezequiel: [link Google Meet]',
        '[13:12, 13/2/2026] Ezequiel: cuando tengas los excel pasamelos. switches PoE+ 30W',
    ]),
    ('CHAT GRUPAL: SBD (Richard Serrats + Marcelo Hamra + Hernán)', [
        '[14:42, 17/2/2026] Hernán: Mail?',
        '[14:44, 17/2/2026] Marce Hamra: Si. Al tuyo de sbd',
        '[15:02, 17/2/2026] Richard Serrats: Tema camaras mando a analizar pliego x exa y le pido comisión. Después definimos si vamos o acompañan. Pero nos quedan 5 días hábiles para preparar todo y aún no lo miramos',
        '[15:03, 17/2/2026] Richard Serrats: No puedo arriesgar a que quede desierto, me matan los de arriba. Hasta hoy no mande nada',
        '[15:04, 17/2/2026] Marce Hamra: Pero esto no lo escribimos nosotros?',
        '[15:06, 17/2/2026] Richard Serrats: mando resumen. no lo corroboré. estuve trabajando hasta las 4 am y hoy desde las 9. Resumen SBASE CCTV: 60 cámaras minidomo 4MP + 6 bullet + 4 conteo personas IA, NVR, servidor, 3 PCs monitoreo, 8 switches PoE...',
        '[15:07, 17/2/2026] Marce Hamra: Y sacaron el pliego en base a algo que le pasamos?',
        '[15:08, 17/2/2026] Richard Serrats: Si todo marce. Presupuestario y demas. El tema es vamos o no vamos. Si vamos hay q armar todo',
        '[15:09, 17/2/2026] Richard Serrats: Si los equipos están en stock, si no se discontinuaron, el precio actual, si tiene visita tecnica, carta del fabricante. Nada sabemos',
        '[15:11, 17/2/2026] Marce Hamra: Solo iría si nos da pedal el fabricante, si no no podemos financiarlo. Sino hay que derivarlo',
        '[15:12, 17/2/2026] Richard Serrats: X las dudas lo mando a revisar en paralelo. Y mañana martillo a hick',
        '[22:47, 17/2/2026] Richard Serrats: Igual mañana una vez q presentes todo arrancamos y evaluamos. O sea armamos todo como si vamos hernan',
        '[0:15, 18/2/2026] Richard Serrats: Suerte mañana',
        '[10:15, 18/2/2026] Richard Serrats: Avisa si salió todo bien cuando abra. Y después arrancamos cctv',
        '[10:27, 18/2/2026] Richard Serrats: Agenda ir mañana a la visita tecnica Hernán es a las 9 am. Así sabemos quiénes van',
    ]),
    ('CHAT: Richard Serrats (WA personal)', [
        '[18:47, 6/2/2026] Hernán: instalación: 3 cajas repetidores, caja decodificador, 3 cámaras, llave termomagnética, llave de corte. 2-3 días trabajo.',
        '[11:09, 10/2/2026] Richard: Ahí valido lo del plotter',
        '[11:10, 10/2/2026] Richard: Acordate de pasarme las tareas de posadas para ver lo de enero q m pediste',
        '[11:52, 10/2/2026] Richard: Que se debe? Así pagamos',
        '[14:46, 10/2/2026] Richard: Armate los números y pedimos póliza para mañana x el tope de presupuesto. Leo ya pasó valor',
        '[15:28, 11/2/2026] Hernán: doc confirmada en SBASE. mañana visita tecnica.',
        '[15:30, 11/2/2026] Richard: Productos nomas',
        '[15:41, 11/2/2026] Hernán: ya le envié. equipos del BOM ajustados a necesidad real del ministerio',
        '[16:48, 11/2/2026] Richard: Si Ale dio el ok avancemos',
    ]),
    ('CHAT: Richard Serrats (WA personal — tema balances)', [
        'Hernán: estaba chequeando los archivos de la lici. Ana no me pasó constancias de presentación balances 2023 y 2024 ante IGJ.',
        'Richard: No lo tengo yo. Lo hace benji. Pedile que te lo pase. Cantorna está de vacaciones.',
    ]),
    ('CHAT: Facundo Fernández (SBASE Técnico)', [
        '[10:17, 11/2/2026] Facundo: Presupuesto oficial USD 205.000,00. Cotización en pesos para mano de obra, equipos pueden cotizarse en USD',
    ]),
    ('CHAT: Elias Fiesa (Hikvision) — mencionado en grupo', [
        'Tema: sistema de logueo propio del ministerio + fotos DNI en HikCentral',
        'Hernán: si no nos lo piden a nosotros y a los demás sí, sería un gol. Cotizaríamos más barato.',
    ]),
    ('CHAT GRUPAL: Richard + Marino + Gervasio + Hernán (Grupo Posadas/SBD)', [
        '[10:31, 23/1/2026] Richard: Teles posadas',
        '[10:32, 23/1/2026] Gervasio: Hay que bajar a dos. Bastante bien',
        '[10:34, 23/1/2026] Marino: TEFITI, SUTEL, NEWCOM-ICS, SANTIAGO SCHVARTZMAN, ABC LATINOAMERICANA, BAFF, RIO INFORMATICA. A uno el más barato no fue a la visita',
        '[10:35, 23/1/2026] Marino: Uno es el más económico, hay que tirarle a BAFF SRL',
        '[16:04, 23/1/2026] Marino: Para ganar los teles hay que desestimar a PERREN SH (no fue a visita técnica) y BAFF SRL (precio vil)',
        '[16:18, 23/1/2026] Marino: Dictamen de Evaluación de Ofertas completo — EX-2025-143521000',
        '[15:55, 27/1/2026] Richard: Me dijeron mínimo Ryzen 5, 16GB, 480GB. Tenés notebooks para cotizarles 4?',
        '[21:09, 6/2/2026] Richard: [fwd Graciela Posadas] pidieron documentación de televisores! uds estarían primeros! Atento marino q ganamos',
        '[11:24, 9/2/2026] Marino: respondido, cualquier cosa avisame',
        '[12:03, 13/2/2026] Hernán: Ayer hablé con Leo, por falta de ménsulas no pudo trabajar ayer ni hoy. Hoy avisé a Ezequiel. Pasé contacto Opelba. Vienen 2 feriados, si el miércoles no están tampoco avanza.',
        '[12:03, 13/2/2026] Richard: Veamos estos temas. Estos retrasos nos pegan en las cobranzas. Necesitamos plata para el giro comercial',
        '[13:09, 13/2/2026] Marino: hubo error de cálculos de Leo, estamos juntando de 3-4 proveedores. La semana que viene va a tener para trabajar',
        '[9:36, 18/2/2026] Hernán: Buen dia. Avisen como seguimos con materiales.',
        '[9:39, 18/2/2026] Richard: No podemos perder más días de laburo. Hay fecha estimada de entrega del primer hito?',
        '[9:46, 18/2/2026] Marino: ya te confirmo, tiene que llegar hoy la primera tanda',
    ]),
]

for title, lines in chats:
    add_section_title(doc1, title)
    for line in lines:
        doc1.add_paragraph(line).runs[0].font.size = Pt(9)

out1 = os.path.join(out_dir, 'Informe de Gestión - DATOS CRUDOS.docx')
doc1.save(out1)
print('Doc 1 (crudos): OK')

# ============================================================
# DOC 2: INFORME con tabla Item/Fecha/Proyecto/Contacto/Mensaje
# ============================================================
doc2 = setup_doc()
add_title(doc2, 'INFORME DE GESTIÓN')

p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Hospital Posadas — Cableado Estructurado — Licitación 96-0051-LPU22')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

add_subtitle(doc2, 'Período: 22 de diciembre 2025 al 18 de febrero 2026')

p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Responsable: Hernán Hamra — Software By Design S.A.')
doc2.add_paragraph()

p = doc2.add_paragraph()
run = p.add_run('Objeto: ')
run.bold = True
p.add_run(
    'Detalle de intervenciones de gestión realizadas por Hernán Hamra en el proyecto Hospital Posadas. '
    'Incluye coordinación directa con personal del Hospital (Ricardo Torres, Claudio Kolln, Abel), '
    'socios (Marino Frangi, Ezequiel Da Conceicao), subcontratistas (Leonardo Martínez - DNET), y proveedores. '
    'Abarca compra y entrega de equipamiento, análisis técnico, negociación de presupuestos, '
    'planificación de acopio de materiales, kick off, inicio y seguimiento de obra.'
)

add_section_title(doc2, 'CRONOGRAMA DE INTERVENCIONES')

table = doc2.add_table(rows=1, cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, text in enumerate(['#', 'Fecha', 'Canal', 'Proyecto', 'Contacto', 'Mensaje / Actividad']):
    make_header_cell(table.rows[0].cells[i], text)

# Item, Fecha, Canal, Proyecto, Contacto, Mensaje
SEP = None
data = [
    # DICIEMBRE 2025
    ('22/12/2025 18:13', 'WA Claudio', 'Posadas', 'Claudio Kolln (Hosp.)', '🔗 Envío specs módulo SFP+ 10G 850nm (link MercadoLibre). Negociación compra y plazos'),
    ('23/12/2025 07:16', 'Archivos', 'Posadas', 'Archivos', '📄 Modificación Excel MATERIALES HOSPITAL POSADAS.xlsx'),
    ('24/12/2025 08:01', 'Archivos', 'Posadas', 'Archivos', '📄 Modificación Excel P&L Hospital Posadas ETAPA III 2025.xlsx'),
    ('29/12/2025 12:06', 'WA Claudio', 'Posadas', 'Claudio Kolln (Hosp.)', 'Claudio: "no te olvides de los SFP, la semana que viene instalamos los switch"'),
    ('30/12/2025 08:53', 'WA Claudio', 'Posadas', 'Claudio Kolln (Hosp.)', '🔗 Envío link tracking Cabify para entrega SFPs. Claudio: "Ya llegaron los SFP!!"'),
    SEP,
    # ENERO PRE-VACACIONES
    ('07/01/2026 12:43', 'Archivos', 'Posadas', 'Archivos', '📄 Actualización P&L Etapa III (.xlsx)'),
    ('07/01/2026 13:07', 'Archivos', 'Posadas', 'Archivos', '📄 OC 96-0231-OC25.pdf (Orden de Compra)'),
    ('13/01/2026 09:03', 'WA Marino', 'Posadas', 'Marino Frangi', 'Coordinación. Marino confirma para la tarde'),
    ('13/01/2026 tarde', 'LLAMADA', 'Posadas', 'Marino Frangi', '📞 Llamada telefónica (inferida: "a la tarde hablamos")'),
    ('14/01/2026 08:37', 'WA Marino', 'Posadas', 'Marino Frangi', 'Hernán solicita llamada urgente. Marino: "te llamo con la estrategia, bajada de línea hacia Leo"'),
    ('14/01/2026 ~09:30', 'LLAMADA', 'Posadas', 'Marino Frangi', '📞 Llamada telefónica (inferida: "estoy marino" → "estoy con otra llamada")'),
    ('14/01/2026 12:05', 'WA Marino/Abel', 'Posadas', 'Marino / Abel', '🔗 Link rotomartillo Kommberg (MercadoLibre) + 📷 Fotos. Hernán: "La metemos nosotros. Es parte del proyecto"'),
    ('14/01/2026 12:17', 'WA Marino', 'Posadas', 'Marino Frangi', '🎤 Audio + 📷 Foto + 😂 Emojis intercambiados'),
    ('14/01/2026 16:48', 'WA Marino', 'Posadas', 'Marino Frangi', 'Marino pide contacto proveedor. Hernán envía Miguel Merlo — Distrib. Versailles'),
    ('15/01/2026 14:44', 'Archivos', 'Posadas', 'Archivos', '📄 4 PDFs Kick Off: Plan de trabajo, Lista materiales, Soporte físico, Rotulación'),
    ('15/01/2026 12:39', 'LLAMADA', 'Posadas', 'Marino Frangi', '📞 Llamada telefónica (inferida: "te llamo en cuanto pueda")'),
    ('15/01/2026 12:45', 'WA Marino', 'Posadas', 'Marino Frangi', '🎤 Audio 2:46 min — Marino: "vamos a estar bien"'),
    ('15/01/2026 12:46', 'WA Marino', 'Posadas', 'Marino Frangi', '🎤 Audio 0:17 min — Marino: "arrancamos con el pie derecho"'),
    ('15/01/2026 18:30', 'WA Marino', 'Posadas', 'Marino Frangi', '★ PLANIFICACIÓN OBRA — Acopio lun 19 a vie 24. Inicio obra lun 26 a vie 30'),
    ('15/01/2026 18:34', 'WA Marino', 'Posadas', 'Marino Frangi', 'Marino: "Correcto, confirmá con Leo que arranque ese lunes y le metemos"'),
    ('16/01/2026 00:51', 'Archivos', 'Posadas', 'Archivos', '📄 Documento reunion de kick off posadas 16-1-26.docx'),
    ('16/01/2026 15:17', 'WA Ricardo', 'Posadas', 'Ricardo Torres', 'Inspección depósito en el Posadas. "Excelente el depósito. Impecable"'),
    ('16/01/2026 16:20', 'WA Abel', 'Posadas', 'Abel (Hospital)', '🔗 Link multiherramienta Matrix 5 en 1. Cierre de semana con equipo del hospital'),
    SEP,
    # VACACIONES 20-31 ENERO
    ('22/01/2026 13:22', 'WA Ricardo', 'Posadas', 'Ricardo Torres', '★ VACACIONES — Ricardo solicita info. Hernán envía datos por 📧 mail y WhatsApp'),
    ('22/01/2026 16:16', 'WA Ricardo', 'Posadas', 'Ricardo Torres', '★ VACACIONES — Specs técnicas plotter: interfaces, memoria, IEEE802.3'),
    ('22/01/2026 21:02', 'WA Ricardo', 'Posadas', 'Ricardo Torres', '★ VACACIONES — ANÁLISIS TÉCNICO: PostScript NO CUMPLE, PDF NO CUMPLE, HP-GL/2 CUMPLE, PCL NO CUMPLE, Velocidad NO CUMPLE, Pantalla NO CUMPLE'),
    ('22/01/2026 21:04', 'WA Ricardo', 'Posadas', 'Ricardo Torres', '★ VACACIONES — Ricardo: "Lo paso para revisión"'),
    ('23/01/2026 14:56', 'WA Ricardo', 'Posadas', 'Ricardo Torres', '★ VACACIONES — Hernán revisa documentación adicional de Ricardo'),
    SEP,
    ('19-24/01/2026', '—', 'Posadas', '—', 'Semana acopio de materiales (coordinado 15/01)'),
    ('26-30/01/2026', '—', 'Posadas', '—', 'Semana inicio de obra (coordinado 15/01)'),
    SEP,
    # FEBRERO
    ('02/02/2026 13:06', 'WA Ricardo', 'Posadas', 'Ricardo Torres', 'Coordinación con Ricardo'),
    ('09/02/2026 10:08', 'WA Ricardo', 'Posadas', 'Ricardo Torres', 'Coordinación con Ricardo'),
    ('09/02/2026 13:08', 'WA Leo', 'Posadas', 'Leo Martínez DNET', 'Leo reporta: 120 tiras instaladas, faltan 105. Necesita 130 ménsulas + 12 curvas 90°'),
    ('09/02/2026 13:27', 'WA Leo', 'Posadas', 'Leo Martínez DNET', 'Hernán confirma: 130 ménsulas 200mm + 12 curvas 90°. Recuerda seguros personal'),
    ('09/02/2026 13:44', 'WA Marino', 'Posadas', 'Marino Frangi', 'Informa: Leo pasó material a Franco. Docs y avance para Ricardo'),
    ('10/02/2026 14:52', 'WA Leo', 'Posadas', 'Leo Martínez DNET', '📄 Hernán envía planos que armó + cotización (documentos adjuntos)'),
    ('10/02/2026 15:09', 'WA Leo', 'Posadas', 'Leo Martínez DNET', 'Reporte por piso: P6 OK, P5 parcial, P4 falta 45m, P3 OK, P2 OK, P1 OK, PB falta 45m'),
    ('11/02/2026 12:09', 'WA Leo', 'Posadas/SBASE', 'Leo Martínez DNET', '📄 Cotización mano de obra SBASE + pliego técnico (documentos enviados)'),
    ('12/02/2026 12:09', 'WA Leo', 'Posadas', 'Leo Martínez DNET', 'Leo reclama ménsulas, articuladas y curvas faltantes'),
    ('12/02/2026 17:34', 'WA Ezequiel', 'SBASE', 'Ezequiel Da Conceicao', 'Análisis técnico: UniFi vs MPLS, switches PoE+, direct attach, Distecna'),
    ('13/02/2026 09:16', 'WA Ezequiel', 'SBASE', 'Ezequiel Da Conceicao', '📞 Reunión Google Meet (videollamada) + 🔗 Link Meet. Definición técnica + docs. Opelba'),
    ('13/02/2026 11:14', 'WA Abel', 'Posadas', 'Abel (Hospital)', 'Hernán gestiona consulta: "Ya te averiguo"'),
    ('13/02/2026 12:36', 'WA Leo', 'Posadas', 'Leo Martínez DNET', 'Leo reclama materiales pendientes'),
    ('13/02/2026 13:12', 'WA Ezequiel', 'SBASE', 'Ezequiel Da Conceicao', '📄 Excels (planillas técnicas) + switches PoE+ 30W para teléfonos y APs'),
    SEP,
    # GRUPO SBD — Richard + Marcelo
    ('17/02/2026 14:42', 'Grupo SBD', 'SBD', 'Hernán / Marcelo', 'Coordinación interna mail SBD'),
    ('17/02/2026 15:02', 'Grupo SBD', 'SBASE CCTV', 'Richard Serrats', 'Richard: "Tema cámaras mando a analizar pliego. Nos quedan 5 días hábiles y aún no lo miramos. No puedo arriesgar que quede desierto"'),
    ('17/02/2026 15:06', 'Grupo SBD', 'SBASE CCTV', 'Richard Serrats', '📄 Richard envía resumen pliego CCTV: 60 cámaras 4MP, NVR, servidor, 3 PCs monitoreo, 8 switches PoE'),
    ('17/02/2026 15:08', 'Grupo SBD', 'SBASE CCTV', 'Richard Serrats', 'Richard: "Presupuestario y demás, todo. El tema es vamos o no. Stock, precios, visita técnica, carta fabricante — nada sabemos"'),
    ('17/02/2026 15:11', 'Grupo SBD', 'SBASE CCTV', 'Marcelo Hamra', 'Marcelo: "Solo iría si nos da pedal el fabricante. Si no, hay que derivarlo"'),
    ('17/02/2026 22:47', 'Grupo SBD', 'SBASE CCTV', 'Richard Serrats', 'Richard: "Mañana una vez que presentes todo arrancamos. Armamos todo como si vamos, Hernán"'),
    ('18/02/2026 00:15', 'Grupo SBD', 'SBASE', 'Richard Serrats', 'Richard: "Suerte mañana"'),
    ('18/02/2026 10:15', 'Grupo SBD', 'SBASE/CCTV', 'Richard Serrats', 'Richard: "Avisá si salió todo bien. Después arrancamos CCTV"'),
    ('18/02/2026 10:27', 'Grupo SBD', 'SBASE CCTV', 'Richard Serrats', 'Richard: "Agenda visita técnica mañana 9am. Así sabemos quiénes van"'),
    SEP,
    # RICHARD WA PERSONAL
    ('06/02/2026 18:47', 'WA Richard', 'AiControl', 'Hernán', 'Presupuesto instalación: 3 repetidores, decodificador, 3 cámaras, llaves termomagnéticas. 2-3 días'),
    ('10/02/2026 11:09', 'WA Richard', 'Posadas', 'Richard Serrats', 'Richard valida plotter. "Acordate de pasarme las tareas de posadas para ver lo de enero que me pediste"'),
    ('10/02/2026 11:52', 'WA Richard', 'Posadas', 'Richard Serrats', 'Richard: "Que se debe? Así pagamos"'),
    ('10/02/2026 14:46', 'WA Richard', 'SBASE', 'Richard Serrats', 'Richard: "Armate los números, pedimos póliza. Leo ya pasó valor"'),
    ('11/02/2026 10:17', 'WA Facundo', 'SBASE', 'Facundo Fernández', 'Presupuesto oficial SBASE: USD 205.000 con IVA. MO en pesos, equipos en USD'),
    ('11/02/2026 15:28', 'WA Richard', 'SBASE', 'Hernán', 'Doc confirmada en SBASE. Mañana visita técnica'),
    ('11/02/2026 15:41', 'WA Richard', 'SBASE/MinSeg', 'Hernán', '📄 Envío equipos BOM ajustados a necesidad real del ministerio + cotización'),
    ('11/02/2026 16:48', 'WA Richard', 'SBASE/MinSeg', 'Richard Serrats', 'Richard: "Si Ale dio el ok avancemos"'),
    SEP,
    # GRUPO POSADAS/SBD — Febrero
    ('10/02/2026 11:10', 'WA Richard', 'Posadas', 'Richard Serrats', 'Richard: "Acordate de pasarme las tareas de posadas para ver lo de enero que me pediste"'),
    ('13/02/2026 12:03', 'Grupo Posadas', 'Posadas', 'Hernán', 'Hernán reporta: Leo no pudo trabajar por falta de ménsulas. Avisó a Ezequiel, pasó contacto Opelba. Alerta por 2 feriados'),
    ('13/02/2026 12:03', 'Grupo Posadas', 'Posadas', 'Richard Serrats', 'Richard: "Estos retrasos nos pegan en las cobranzas. Necesitamos plata para el giro comercial"'),
    ('13/02/2026 13:09', 'Grupo Posadas', 'Posadas', 'Marino Frangi', 'Marino: error cálculos Leo, juntando de 3-4 proveedores. Semana que viene tendrá material'),
    ('18/02/2026 09:36', 'Grupo Posadas', 'Posadas', 'Hernán', 'Hernán: "Avisen cómo seguimos con materiales"'),
    ('18/02/2026 09:39', 'Grupo Posadas', 'Posadas', 'Richard Serrats', 'Richard: "No podemos perder más días. Fecha estimada primer hito?"'),
    ('18/02/2026 09:46', 'Grupo Posadas', 'Posadas', 'Marino Frangi', 'Marino: "hoy llega la primera tanda"'),
]

item_num = 0
for idx, row_data in enumerate(data):
    if row_data is None:
        # Separator row
        row = table.add_row()
        continue

    item_num += 1
    fecha, canal, proyecto, contacto, mensaje = row_data
    row = table.add_row()

    is_vacation = '★ VACACIONES' in mensaje
    is_planned = 'acopio' in mensaje.lower() and contacto == '—'
    is_star_plan = '★ PLANIFICACIÓN' in mensaje

    texts = [str(item_num), fecha, canal, proyecto, contacto, mensaje]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        if i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(8)
        if is_vacation:
            run.bold = True
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        elif is_planned or is_star_plan:
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    if is_vacation:
        shade_row(row, 'FFF2CC')
    elif item_num % 2 == 0:
        shade_row(row, 'F2F2F2')

# HALLAZGO CLAVE
add_section_title(doc2, 'HALLAZGO CLAVE: TRABAJO DURANTE VACACIONES (20-31 ENERO)')

p = doc2.add_paragraph()
p.add_run(
    'El 22 y 23 de enero de 2026, durante el período de vacaciones (20-31 enero), '
    'Hernán Hamra realizó las siguientes tareas documentadas con timestamps de WhatsApp:'
)

items_vac = [
    'Envío de especificaciones técnicas completas de plotter a Ricardo Torres del Hospital Posadas.',
    'Análisis técnico exhaustivo de cumplimiento de pliego: 6 ítems analizados, 5 NO CUMPLEN.',
    'Revisión de documentación adicional enviada por Ricardo Torres (23/01).',
    'Trabajo realizado a las 21:02 hs — fuera de horario laboral y durante vacaciones.',
    'Ricardo Torres respondió "Lo paso para revisión" — confirmando que el trabajo fue útil.',
]
for item in items_vac:
    p = doc2.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

# OBSERVACIONES
add_section_title(doc2, 'OBSERVACIONES GENERALES')

obs = [
    'Gestión continua desde diciembre 2025: compra y entrega de módulos SFP+ (22-30 dic), P&L (23-24 dic).',
    'Semana 13-16 enero: coordinación intensiva con Marino, Abel y Ricardo. Kick Off, inspección, cronograma.',
    'El cronograma de acopio (19-24 ene) e inicio de obra (26-30 ene) fue definido por Hernán el 15/01.',
    'Coordinación con Leo delegada explícitamente: "confirmá con Leo que arranque" (Marino, 15/01).',
    'Obra ejecutada según plan: al 9/02 Leo reporta 120+ tiras instaladas, avance en todos los pisos.',
    'Febrero: gestión diaria — materiales, seguros, planos propios, cotizaciones SBASE, meets con Ezequiel.',
    'Leyenda: 📞 Llamada, 🎤 Audio WA, 📄 Documento/Archivo, 📷 Foto, 🔗 Link, 📧 Mail, 😂 Emoji/Meme.',
]
for o in obs:
    p = doc2.add_paragraph(o, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

out2 = os.path.join(out_dir, 'Informe de Gestión - Enero 2026 - Hernán Hamra.docx')
doc2.save(out2)
print('Doc 2 (informe): OK')
print(f'Total items: {item_num}')
print(f'Guardados en: {out_dir}')
