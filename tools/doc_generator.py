"""
ARGOS - Generador de documentos Word
Crea documentos .docx desde template preservando header/footer.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Colores corporativos SBD
COLOR_AZUL = (27, 79, 114)
COLOR_GRIS = (93, 109, 126)


def load_template(template_path):
    """Carga un template y limpia el body preservando sectPr (header/footer)."""
    doc = Document(template_path)
    body = doc.element.body
    for child in list(body):
        if child.tag != '{%s}sectPr' % NS:
            body.remove(child)
    return doc


def add_para(doc, text, bold=False, size=11, align=None, space_after=6,
             space_before=2, font='Calibri', color=None):
    """Agrega un párrafo con formato completo incluyendo rFonts via lxml."""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = font
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        # Forzar font via lxml (necesario para que Word respete el font)
        r_elem = run._element
        rPr = r_elem.find('{%s}rPr' % NS)
        if rPr is None:
            rPr = etree.SubElement(r_elem, '{%s}rPr' % NS)
        rFonts = rPr.find('{%s}rFonts' % NS)
        if rFonts is None:
            rFonts = etree.SubElement(rPr, '{%s}rFonts' % NS)
        rFonts.set('{%s}ascii' % NS, font)
        rFonts.set('{%s}hAnsi' % NS, font)
        rFonts.set('{%s}cs' % NS, font)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def add_separator(doc):
    """Agrega una línea separadora centrada en azul."""
    p = doc.add_paragraph()
    run = p.add_run('_' * 60)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*COLOR_AZUL)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(12)


def add_firma(doc, nombre, cargo, empresa, extra=None):
    """Agrega bloque de firma estándar."""
    add_para(doc, '', space_after=24)
    add_para(doc, '_________________________________', size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, nombre, size=11, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, cargo, size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, empresa, size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    if extra:
        add_para(doc, extra, size=11,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
